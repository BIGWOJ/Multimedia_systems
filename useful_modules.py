import numpy as np
import matplotlib.pyplot as plt

def generate_figures(signal, fs, time_margin=[0, 0.02]):
    import scipy.fftpack

    time = np.arange(0, len(signal)) / fs

    plt.figure()
    plt.subplot(2, 1, 1)
    plt.plot(time, signal)
    plt.xlabel('[s]')
    plt.ylabel('Amplitude')
    plt.xlim(time_margin)
    plt.title('Signal')

    yf = scipy.fftpack.fft(signal)
    xf = np.linspace(0, fs / 2, len(yf) // 2)

    plt.subplot(2, 1, 2)
    plt.plot(xf, 20 * np.log10(np.abs(yf[:len(yf) // 2])))
    plt.xlabel('[Hz]')
    plt.ylabel('[dB]')
    plt.title('1/2 Spectrum')
    plt.tight_layout()
    plt.show()

def generate_docx_template():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert

    document = Document()
    document.add_heading('Change title', 0)

    files = ['sin60Hz.wav', 'sin440Hz.wav', 'sin8000Hz.wav']
    margins = [[0, 0.02], [0.133, 0.155]]

    for file in files:
        document.add_heading('Plik - {}'.format(file), 2)
        for i, Margin in enumerate(margins):
            document.add_heading('Time margin {}'.format(Margin), 3)
            fig, axs = plt.subplots(2, 1, figsize=(10, 7))

            ############################################################
            # Analizing code, plots
            ############################################################

            fig.suptitle('Time margin {}'.format(Margin))
            fig.tight_layout(pad=1.5)
            memfile = BytesIO()
            fig.savefig(memfile)

            document.add_picture(memfile, width=Inches(6))

            memfile.close()
            ############################################################
            # Text data - values, outputs etc
            document.add_paragraph('Example value = {}'.format(np.random.rand(1)))
            ############################################################

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')