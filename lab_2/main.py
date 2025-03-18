import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd

def task_1():
    img_file = 'IMG_INTRO/A3.png'
    img_imread = plt.imread(img_file)

    def img_to_uint8(img):
        if np.issubdtype(img.dtype, np.unsignedinteger):
            return img
        else:
            return (img * 255).astype('uint8')

    def img_to_float(img):
        if np.issubdtype(img.dtype, np.floating):
            return img
        else:
            return img / 255

    plt.imshow(img_imread)
    plt.show()

    R_color = img_imread[:,:,0]
    G_color = img_imread[:,:,1]
    B_color = img_imread[:,:,2]

    plt.imshow(R_color, cmap=plt.cm.gray)
    plt.show()

    img_gray = 0.2126 * R_color + 0.7152 * G_color + 0.0722 * B_color

    img_cv = cv2.imread(img_file)
    plt.imshow(img_cv)
    plt.show()

    img_cv_rgb = cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB)
    plt.imshow(img_cv_rgb)
    plt.show()

    img_cv_bgr = cv2.cvtColor(img_cv_rgb, cv2.COLOR_RGB2BGR)
    plt.imshow(img_cv_bgr)
    plt.show()

def task_2(img, task_3=False):
    fig = plt.figure(figsize=(10, 10))
    plt.subplot(3,3,1)
    plt.imshow(img)

    R_color = img[:,:,0]
    plt.subplot(3,3,2)
    plt.imshow(R_color, cmap=plt.cm.gray)

    G_color = img[:,:,1]
    B_color = img[:,:,2]
    img_y2 = 0.2126 * R_color + 0.7152 * G_color + 0.0722 * B_color
    plt.subplot(3,3,3)
    plt.imshow(img_y2, cmap=plt.cm.gray)

    plt.subplot(3,3,4)
    plt.imshow(R_color)

    plt.subplot(3, 3, 5)
    plt.imshow(G_color)

    plt.subplot(3, 3, 6)
    plt.imshow(B_color)

    plt.subplot(3, 3, 7)
    img_copy = img.copy()
    img_copy[:, :, 1] = 0
    img_copy[:, :, 2] = 0
    plt.imshow(img_copy)

    plt.subplot(3, 3, 8)
    img_copy = img.copy()
    img_copy[:, :, 0] = 0
    img_copy[:, :, 2] = 0
    plt.imshow(img_copy)

    plt.subplot(3, 3, 9)
    img_copy = img.copy()
    img_copy[:, :, 0] = 0
    img_copy[:, :, 1] = 0
    plt.imshow(img_copy)
    plt.tight_layout()

    if task_3:
        return fig
    else:
        plt.show()

def task_3():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO

    df = pd.DataFrame(data={'Filename': ['IMG_INTRO/B02.jpg'], 'Grayscale': [False],
                            'Fragments': [[[0,0,800,1200], [0,0,200,200], [200,0,400,200], [400,0,600,200], [600,0,800,200]]]
                            })

    document = Document()
    document.add_heading('Analiza fragmentów obrazu', 0)

    for index, row in df.iterrows():
        document.add_heading('Obraz {}'.format(row['Filename']), 3)
        img = cv2.imread(row['Filename'])
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        for i, f in enumerate(row['Fragments']):
            fragment = img[f[0]:f[2], f[1]:f[3]].copy()
            fig = task_2(fragment, task_3=True)

            memfile = BytesIO()
            fig.savefig(memfile)
            memfile.seek(0)

            document.add_heading(f'Fragment {i + 1}', level=3)
            document.add_picture(memfile, width=Inches(6))

            memfile.close()
            plt.close(fig)

        document.save('report.docx')

img = plt.imread('IMG_INTRO/B02.jpg')
# task_1()
# task_2(img)
task_3()