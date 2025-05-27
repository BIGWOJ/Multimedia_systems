import cv2
import matplotlib.pyplot as plt
import numpy as np
from skimage.metrics import structural_similarity
from PIL import Image

def water_mark(img,mask,alpha=0.25):
    assert (img.shape[0]==mask.shape[0]) and (img.shape[1]==mask.shape[1]), "Wrong size"
    if len(img.shape)<3:
        flag=True
        t_img=cv2.cvtColor(img,cv2.COLOR_GRAY2RGBA)
    else:
        flag=False
        t_img=cv2.cvtColor(img,cv2.COLOR_RGB2RGBA)
    if (mask.dtype==bool):
        t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
    elif (mask.dtype==np.uint8):
        if len(mask.shape)<3:
            t_mask=cv2.cvtColor((mask).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
        else:
            t_mask=cv2.cvtColor((mask).astype(np.uint8),cv2.COLOR_RGB2RGBA)
    else:
        if len(mask.shape)<3:
            t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_GRAY2RGBA)
        else:
            t_mask=cv2.cvtColor((mask*255).astype(np.uint8),cv2.COLOR_RGB2RGBA)
    t_out=cv2.addWeighted(t_img,1,t_mask,alpha,0)
    if flag:
        out=cv2.cvtColor(t_out,cv2.COLOR_RGBA2GRAY)
    else:
        out=cv2.cvtColor(t_out,cv2.COLOR_RGBA2RGB)
    return out

def put_data(img,data,binary_mask=np.uint8(1)):
    assert img.dtype==np.uint8 , "img wrong data type"
    assert binary_mask.dtype==np.uint8, "binary_mask wrong data type"
    un_binary_mask=np.unpackbits(binary_mask)
    if data.dtype!=bool:
        unpacked_data=np.unpackbits(data)
    else:
        unpacked_data=data
    dataspace=img.shape[0]*img.shape[1]*np.sum(un_binary_mask)
    assert (dataspace>=unpacked_data.size) , "too much data"
    if dataspace==unpacked_data.size:
        prepered_data=unpacked_data.reshape(img.shape[0],img.shape[1],np.sum(un_binary_mask)).astype(np.uint8)
    else:
        prepered_data=np.resize(unpacked_data,(int(img.shape[0]),int(img.shape[1]),int(np.sum(un_binary_mask)))).astype(np.uint8)
    mask=np.full((img.shape[0],img.shape[1]),binary_mask)
    img=np.bitwise_and(img,np.invert(mask))
    bv=0
    for i,b in enumerate(un_binary_mask[::-1]):
        if b:
            temp=prepered_data[:,:,bv]
            temp=np.left_shift(temp,i)
            img=np.bitwise_or(img,temp)
            bv+=1
    return img

def pop_data(img,binary_mask=np.uint8(1),out_shape=None):
    un_binary_mask=np.unpackbits(binary_mask)
    data=np.zeros((img.shape[0],img.shape[1],np.sum(un_binary_mask))).astype(np.uint8)
    bv=0
    for i,b in enumerate(un_binary_mask[::-1]):
        if b:
            mask=np.full((img.shape[0],img.shape[1]),2**i)
            temp=np.bitwise_and(img,mask)
            data[:,:,bv]=temp[:,:].astype(np.uint8)
            bv+=1
    if out_shape!=None:
        tmp=np.packbits(data.flatten())
        tmp=tmp[:np.prod(out_shape)]
        data=tmp.reshape(out_shape)
    return data

def calculate_metrics(original, modified):
    mse = np.mean((original - modified) ** 2)
    psnr = 10 * np.log10((255 ** 2) / mse)
    if original.ndim == 1:
        original = original.reshape(1, -1)
    if modified.ndim == 1:
        modified = modified.reshape(1, -1)
    ssim = structural_similarity(original, modified, multichannel=(original.ndim == 3),
                                 channel_axis=2 if original.ndim == 3 else None)
    return psnr, ssim

def task_1(img, generate_report=False):
    blue_channel = img[:, :, 0]
    with open("text.txt", "r") as file:
        text = file.read()

    binary_mask = np.uint8(1)
    text_data = np.frombuffer(text.encode('utf-8'), dtype=np.uint8)
    encoded = put_data(blue_channel, text_data, binary_mask)
    encoded_img = img.copy()
    encoded_img[:, :, 0] = encoded
    restored_bits = pop_data(encoded, binary_mask, out_shape=text_data.shape)
    restored_text = restored_bits.tobytes().decode('utf-8')

    fig = plt.figure(figsize=(10, 5))
    metrics = calculate_metrics(encoded_img[:, :, 0], encoded_img[:, :, 0])
    plt.subplot(1, 2, 1)
    plt.imshow(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
    plt.title(f'Carrier image\nB channel: {get_metrics_text(metrics)}')

    metrics = calculate_metrics(blue_channel, encoded_img[:, :, 0])
    plt.subplot(1, 2, 2)
    plt.imshow(cv2.cvtColor(encoded_img, cv2.COLOR_BGR2RGB))
    plt.title(f'Carrier image with hidden text\nB channel: {get_metrics_text(metrics)}')

    if generate_report:
        return text, restored_text, encoded_img, fig

    plt.show()
    print(text,'\n\n',restored_text)

def get_metrics_text(metrics):
    return f"PSNR: {metrics[0]:.2f}, SSIM: {metrics[1]:.2f}"

def task_2_3(carrier_path, img_hidden_path, generate_report=False, only_task_2=False):
    def recover_image(stego_img, shape, RG_mask, B_mask, dtype=np.uint8):
        recovered = np.zeros(shape, dtype=dtype)

        masks = {
            'R': RG_mask,
            'G': RG_mask,
            'B': B_mask,
        }

        for channel_idx, (channel_name, mask) in enumerate(masks.items()):
            recovered[:, :, channel_idx] = pop_data(stego_img[:, :, channel_idx], mask, out_shape=shape[:2])

        return recovered

    img_carrier = cv2.cvtColor(cv2.imread(carrier_path), cv2.COLOR_BGR2RGB)
    img_hidden = cv2.cvtColor(cv2.imread(img_hidden_path), cv2.COLOR_BGR2RGB)

    R = img_carrier[:, :, 0].copy()
    G = img_carrier[:, :, 1].copy()
    B = img_carrier[:, :, 2].copy()
    R2 = img_hidden[:, :, 0].copy()
    G2 = img_hidden[:, :, 1].copy()
    B2 = img_hidden[:, :, 2].copy()

    carrier_with_hidden = np.copy(img_carrier)
    binary_mask_rg = np.uint8(0b00000011)
    binary_mask_b = np.uint8(0b00000111)
    carrier_with_hidden[:, :, 0] = put_data(R, R2, binary_mask_rg)
    carrier_with_hidden[:, :, 1] = put_data(G, G2, binary_mask_rg)
    carrier_with_hidden[:, :, 2] = put_data(B, B2, binary_mask_b)

    recovered_img = recover_image(carrier_with_hidden, img_hidden.shape, binary_mask_rg, binary_mask_b, dtype=img_hidden.dtype)

    fig = plt.figure(figsize=(10, 12))

    plt.subplot(3,4,1)
    plt.imshow(img_carrier)
    metrics_b = calculate_metrics(img_carrier[:, :, 2], img_carrier[:, :, 2])
    plt.title(f'Carrier image\n B channel: {get_metrics_text(metrics_b)}')
    plt.axis('off')

    plt.subplot(3,4,2)
    plt.imshow(img_carrier[:, :, 0], cmap='Reds')
    plt.title('R channel')
    plt.axis('off')

    plt.subplot(3, 4,3)
    plt.imshow(img_carrier[:, :, 1], cmap='Greens')
    plt.title('G channel')
    plt.axis('off')

    plt.subplot(3, 4,4)
    plt.imshow(img_carrier[:, :, 2], cmap='Blues')
    plt.title('B channel')
    plt.axis('off')

    plt.subplot(3,4,5)
    plt.imshow(carrier_with_hidden)
    metrics_b = calculate_metrics(img_carrier[:, :, 2], carrier_with_hidden[:, :, 0])
    plt.title(f'Carrier with hidden image\n B channel: {get_metrics_text(metrics_b)}')
    plt.axis('off')

    plt.subplot(3,4,6)
    plt.imshow(carrier_with_hidden[:, :, 0], cmap='Reds')
    plt.title('R channel')
    plt.axis('off')

    plt.subplot(3,4,7)
    plt.imshow(carrier_with_hidden[:, :, 1], cmap='Greens')
    plt.title('G channel')
    plt.axis('off')

    plt.subplot(3,4,8)
    plt.imshow(carrier_with_hidden[:, :, 2], cmap='Blues')
    plt.title('B channel')
    plt.axis('off')

    plt.subplot(3,4,9)
    plt.imshow(recovered_img)
    plt.title('Hidden image')
    plt.subplots_adjust(hspace=0.5)
    plt.axis('off')

    if not generate_report:
        plt.show()

    if only_task_2:
        return fig

    ########## Task 3 ##########

    RG_masks = [np.uint8(0b00000011), np.uint8(0b00000111), np.uint8(0b00001111), np.uint8(0b00011111), np.uint8(0b00111111), np.uint8(0b01111111), np.uint8(0b11111111)]
    B_masks = [np.uint8(0b00000111)] + [mask for mask in RG_masks[2:]]
    figs = []
    for RG_mask in RG_masks:
        for B_mask in B_masks:
            img_carrier = cv2.cvtColor(cv2.imread(carrier_path), cv2.COLOR_BGR2RGB)
            img_hidden = cv2.cvtColor(cv2.imread(img_hidden_path), cv2.COLOR_BGR2RGB)
            fig = plt.figure(figsize=(10, 12))
            R = img_carrier[:, :, 0].copy()
            G = img_carrier[:, :, 1].copy()
            B = img_carrier[:, :, 2].copy()
            R2 = img_hidden[:, :, 0].copy()
            G2 = img_hidden[:, :, 1].copy()
            B2 = img_hidden[:, :, 2].copy()

            carrier_with_hidden = np.copy(img_carrier)
            carrier_with_hidden[:, :, 0] = put_data(R, R2, RG_mask)
            carrier_with_hidden[:, :, 1] = put_data(G, G2, RG_mask)
            carrier_with_hidden[:, :, 2] = put_data(B, B2, B_mask)

            recovered_img = recover_image(carrier_with_hidden, img_hidden.shape, RG_mask, B_mask, img_hidden.dtype)

            plt.subplot(3, 4, 1)
            plt.suptitle(f'RG mask {RG_mask:#010b}, B mask {B_mask:#010b}')
            plt.imshow(img_carrier)
            plt.title('Carrier image')
            plt.axis('off')

            plt.subplot(3, 4, 2)
            plt.imshow(img_carrier[:, :, 0], cmap='Reds')
            plt.title('R channel')
            plt.axis('off')

            plt.subplot(3, 4, 3)
            plt.imshow(img_carrier[:, :, 1], cmap='Greens')
            plt.title('G channel')
            plt.axis('off')

            plt.subplot(3, 4, 4)
            plt.imshow(img_carrier[:, :, 2], cmap='Blues')
            plt.title('B channel')
            plt.axis('off')

            plt.subplot(3, 4, 5)
            plt.imshow(carrier_with_hidden)
            metrics_r = calculate_metrics(img_carrier[:, :, 0], carrier_with_hidden[:, :, 0])
            metrics_g = calculate_metrics(img_carrier[:, :, 1], carrier_with_hidden[:, :, 0])
            metrics_b = calculate_metrics(img_carrier[:, :, 2], carrier_with_hidden[:, :, 0])
            plt.title(f'Carrier with hidden image\n'
                      f'R channel: {get_metrics_text(metrics_b)}\n'
                      f'G channel: {get_metrics_text(metrics_g)}\n'
                      f'B channel: {get_metrics_text(metrics_r)}')
            plt.axis('off')

            plt.subplot(3, 4, 6)
            plt.imshow(carrier_with_hidden[:, :, 0], cmap='Reds')
            plt.title('R channel')
            plt.axis('off')

            plt.subplot(3, 4, 7)
            plt.imshow(carrier_with_hidden[:, :, 1], cmap='Greens')
            plt.title('G channel')
            plt.axis('off')

            plt.subplot(3, 4, 8)
            plt.imshow(carrier_with_hidden[:, :, 2], cmap='Blues')
            plt.title('B channel')
            plt.axis('off')

            plt.subplot(3, 4, 9)
            plt.imshow(recovered_img)
            plt.title('Hidden image')
            plt.subplots_adjust(hspace=0.5)
            plt.axis('off')

            if generate_report:
                figs.append(fig)
            else:
                plt.suptitle(f'RG mask {RG_mask:#010b}, B mask {B_mask:#010b}')
                plt.show()
            plt.close(fig)

    if generate_report:
        return figs

def task_4(img, generate_report=False):
    mask = cv2.imread('images/binary.png', 0).astype(np.uint8)
    alpha_values = [0.1, 0.25, 0.5]
    fig = plt.figure(figsize=(10, 8))
    for img_counter, alpha in enumerate(alpha_values):
        watermarked = water_mark(img, mask, alpha=alpha)
        watermarked = cv2.cvtColor(watermarked, cv2.COLOR_BGR2RGB)
        psnr, ssim = calculate_metrics(img, watermarked)

        plt.subplot(3, 1, img_counter + 1)
        plt.subplots_adjust(hspace=0.75)
        plt.imshow(watermarked)
        plt.title(f'Alpha: {alpha}, PSNR: {psnr:.2f}, SSIM: {ssim:.4f}')

    if generate_report:
        return fig
    plt.show()

def add_image_to_docx(document, image, title):
    from io import BytesIO
    from matplotlib.figure import Figure
    from docx.shared import Inches

    memfile = BytesIO()
    if isinstance(image, Figure) or isinstance(image, plt.Axes):
        image.savefig(memfile)
        plt.close(image)
    else:
        fig = plt.figure()
        plt.imshow(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
        plt.title(title)
        fig.savefig(memfile)
        plt.close(fig)
    memfile.seek(0)
    document.add_picture(memfile, width=Inches(7))
    memfile.close()

def generate_report(img_path, img_hidden_path):
    from docx import Document
    from docx2pdf import convert
    from os import remove

    document = Document()
    document.add_heading('Steganografia i znaki wodne\nWojciech Latos', 0)

    document.add_heading('Zadanie 1', level=1)
    text, restored_text, encoded_img, fig = task_1(cv2.imread(img_path), True)
    add_image_to_docx(document, fig, '')
    document.add_paragraph(f'Tekst ukryty w obrazie: {text}')
    document.add_paragraph(f'Odzyskany tekst: {restored_text}')
    document.add_paragraph('Jak widać, tekst został poprawnie ukryty oraz odzyskany. Obrazy nie różnią się wizualnie względem siebie.'
                           'Wartość SSIM pozostała na poziomie 1, co oznacza, że obrazy są identyczne pod względem strukturalnym.\n'
                           'Natomiast wartość PSNR jest bardzo wysoka, co sugeruje, że różnice między obrazami są minimalne i nie są dostrzegalne gołym okiem.')


    document.add_heading('Zadanie 2', level=1)
    fig = task_2_3(img_path, img_hidden_path, generate_report=True, only_task_2=True)
    add_image_to_docx(document, fig, '')

    document.add_heading('Zadanie 3', level=1)
    document.add_paragraph('Na podstawie poniższych zestawień obrazów, minimalny budżet bitowy, przy którym degradacja obrazu nosiciela staje się zauważalna gołym okiem oceniam na wartość 8.\n'
                           'Dla mnie obrazem, który zaczął być zauważalnie zmieniony to ten, dla którego maska wartsw RG była dla 2 bitów - 0b00000011, a dla warstwy B 6 bitów - 0b00111111.\n'
                           'Wartości PSNR oraz SSIM dla kolejnych wartsw RGB były następujące:\n'
                           '27.53, 28.09, 44.19 oraz 0.82, 0.92, 0.98.')
    figs = task_2_3(img_path, img_hidden_path, generate_report=True, only_task_2=False)

    for fig in figs:
        add_image_to_docx(document, fig, '')

    document.add_heading('Zadanie 4', level=1)
    fig = task_4(cv2.imread(img_path), generate_report=True)
    add_image_to_docx(document, fig, '')
    document.add_paragraph('Jak widać na powyższych zdjęciach wraz ze znakami wodnymi, wartość PSNR wzrasta wraz ze wzrostem wartości alpha.\n'
                           'Natomiast SSIM maleje wraz ze wzrostem wartości alpha.\n'
                           'Oczywiście - im większa wartość alpha tym znak jest bardziej widoczny.')


    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')
    remove(docx_path)

img_path = 'images/img_1.png'
img_hidden_path = 'images/img_2.png'
img = cv2.imread(img_path)
img_hidden = cv2.imread(img_hidden_path)

# task_1(img,False)
# task_2_3(img_path, img_hidden_path, generate_report=False, only_task_2=False)
# task_4(img)
generate_report(img_path, img_hidden_path)
