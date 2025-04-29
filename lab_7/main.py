import numpy as np
import matplotlib.pyplot as plt


def A_law_encode(data):
    A = 87.6
    indexes = np.abs(data) < 1/A
    denominator = 1 + np.log(A)
    data[indexes] = np.sign(data[indexes]) * (A * np.abs(data[indexes]) / denominator)
    data[~indexes] = np.sign(data[~indexes]) * ((1 + np.log(A * np.abs(data[~indexes]))) / denominator)

    return data

def A_law_decode(encoded):
    A = 87.6
    indexes = np.abs(encoded) < 1 / (1 + np.log(A))
    encoded[indexes] = np.sign(encoded[indexes]) * (np.abs(encoded[indexes]) * (1 + np.log(A)) / A)
    encoded[~indexes] = np.sign(encoded[~indexes]) * (np.exp(np.abs(encoded[~indexes]) * (1 + np.log(A)) - 1) / A)

    return encoded

def mu_law_encode(data):
    mu = 255
    indexes = (-1 <= data) & (data <= 1)
    data[indexes] = np.sign(data[indexes]) * (np.log(1 + mu * np.abs(data[indexes])) / (np.log(1 + mu)))

    return data

def mu_law_decode(encoded):
    mu = 255
    indexes = (-1 <= encoded) & (encoded <= 1)
    encoded[indexes] = np.sign(encoded[indexes]) * (1 / mu) * ((1 + mu) ** np.abs(encoded[indexes]) -1)

    return encoded

def DPCM_encode(data, bit):
    y = np.zeros(data.shape)
    e = 0
    for i in range(0, data.shape[0]):
        y[i] = quantization(data[i] - e, bit)
        e += y[i]
    return y

def DPCM_decode(encoded):
    decoded = np.zeros(encoded.shape)
    e = 0
    for i in range(encoded.shape[0]):
        decoded[i] = e + encoded[i]
        e = decoded[i]

    return decoded

def DPCM_encode_prediction(x, bit, n):
    encoded = np.zeros(x.shape)
    reconstructed_signal = np.zeros(x.shape)
    e = 0
    for i in range(0, x.shape[0]):
        encoded[i] = quantization(x[i] - e, bit)
        reconstructed_signal[i] = encoded[i] + e
        if i > 0:
            e = np.mean(reconstructed_signal[max(0, i - n):i])
        else:
            e = 0  # No past samples available for the first sample

    return encoded

def DPCM_decode_prediction(encoded, n):
    decoded = np.zeros(encoded.shape)  # Decoded signal
    reconstructed_signal = np.zeros(encoded.shape)  # Reconstructed signal

    e = 0  # Initial prediction value
    for i in range(encoded.shape[0]):
        # Reconstruct the current sample
        decoded[i] = e + encoded[i]

        # Update the reconstructed signal
        reconstructed_signal[i] = decoded[i]

        # Predict the next sample using the mean of the last `n` samples
        if i > 0:
            e = np.mean(reconstructed_signal[max(0, i - n):i])  # Mean of past samples
        else:
            e = 0  # No past samples available for the first sample

    return decoded

def quantization(data, bit):
    levels = 2 ** bit
    step = 2 / (levels - 1)
    quantized = np.round((data + 1) / step) * step - 1

    return quantized

def a():
    x = np.linspace(-1, 1 , 1000)

    x_q = quantization(x, 8)

    y_1 = A_law_encode(x.copy())
    y_1_q = quantization(y_1, 8)

    y_2 = mu_law_encode(x.copy())
    y_2_q = quantization(y_2, 8)

    y_de_a = A_law_decode(y_1.copy())
    y_de_a_q = A_law_decode(y_1_q.copy())

    y_de_mu = mu_law_decode(y_2.copy())
    y_de_mu_q = mu_law_decode(y_2_q.copy())

    fig, axs = plt.subplots(1, 2, figsize=(14, 10))


    axs[0].plot(x, y_1_q, label='Sygnal po kompresji a-law po kwantyzacji do 8-bitow')
    axs[0].plot(x, y_1, label='Sygnal po kompresji a-law bez kwantyzacji')
    axs[0].plot(x, y_2_q, label='Sygnal po kompresji mu-law po kwantyzacji do 8-bitow')
    axs[0].plot(x, y_2, label='Sygnal po kompresji mu-law bez kwantyzacji')
    axs[0].set_title("Krzywa kompresji")
    axs[0].set_xlabel("Oryginalny sygnał")
    axs[0].set_ylabel("Po kompresji")
    axs[0].legend()
    axs[0].grid(True)

    axs[1].plot(x, x, label='Sygnał oryginalny')
    axs[1].plot(x, y_de_a_q, label='Sygnał po dekompresji a-law (quantization 8-bitów)')
    axs[1].plot(x, y_de_mu_q, label='Sygnał po dekompresji mu-law (quantization 8-bitów)')
    axs[1].plot(x, x_q, label='Sygnał oryginalny po kwantyzacji')

    axs[1].set_title("Krzywa dekompresji")
    axs[1].set_xlabel("Oryginalny sygnał")
    axs[1].set_ylabel("Po dekompresji")
    axs[1].legend()
    axs[1].grid(True)



    plt.tight_layout()
    plt.show()

def a_2():
    x = np.linspace(-0.5, -0.25, 1000)
    signal = 0.9 * np.sin(np.pi * x * 4)
    plt.plot(x, signal)
    plt.show()
    a_decoded = A_law_decode(quantization(A_law_encode(signal.copy()), 6))
    plt.plot(x, signal)
    plt.show()
    # exit()
    mu_decoded = mu_law_decode(quantization(mu_law_encode(signal.copy()), 6))
    dpcm_y = DPCM_encode(signal, 6)
    dpcm_decoded = DPCM_decode(dpcm_y)
    dpcm_pred_y = DPCM_encode_prediction(signal, 6, n=3)
    dpcm_pred_decoded = DPCM_decode_prediction(dpcm_pred_y, n=3)

    fig, axs = plt.subplots(5, 1, figsize=(8, 10), sharex=True)
    fig.suptitle("Przykład A kwantyzacja do 6 bitów")

    axs[0].plot(x, signal)
    axs[0].set_title("Oryginalny sygnał")

    axs[1].plot(x, a_decoded)
    axs[1].set_title("Kompresja A-law")

    axs[2].plot(x, mu_decoded)
    axs[2].set_title("Kompresja mu-law")

    axs[3].plot(x, dpcm_decoded)
    axs[3].set_title("Kompresja DPCM bez predykcji")

    axs[4].plot(x, dpcm_pred_decoded)
    axs[4].set_title("Kompresja DPCM z predykcją")

    plt.tight_layout()
    plt.show()

def b():
    x = np.linspace(-0.5, -0.25, 1000)
    signal = 0.9 * np.sin(np.pi * x * 4)

    a_decoded = A_law_decode(quantization(A_law_encode(signal.copy()), 6))
    mu_decoded = mu_law_decode(quantization(mu_law_encode(signal.copy()), 6))
    dpcm_y = DPCM_encode(signal, 6)
    dpcm_decoded = DPCM_decode(dpcm_y)
    dpcm_pred_y = DPCM_encode_prediction(signal.copy(), 6, n=3)
    dpcm_pred_decoded = DPCM_decode_prediction(dpcm_pred_y, n=3)

    plt.figure(figsize=(12, 6))
    plt.plot(x, signal, label='Sygnał oryginalny')
    plt.plot(x, a_decoded, label='Sygnał po dekompresji z a-law')
    plt.plot(x, mu_decoded, label='Sygnał po dekompresji z mu-law')
    plt.plot(x, dpcm_decoded, label='Sygnał po dekompresji z DPCM bez predykcji')
    plt.plot(x, dpcm_pred_decoded, label='Sygnał po dekompresji z DPCM z predykcją')

    plt.title("Porównanie metod dekompresji – Przykład B")
    plt.xlabel("Czas / próbki")
    plt.ylabel("Amplituda")
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.show()

def generate_report():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert
    from os import remove
    import matplotlib.pyplot as plt

    x_ranges = [np.linspace(-1, 1, 1000), np.linspace(-0.9, -0.8, 1000), np.linspace(-0.01, 0.01, 1000)]

    document = Document()
    document.add_heading('Kompresja stratna audio\nWojciech Latos', 0)

    for x_range in x_ranges:
        encoded = A_law_encode(x_range.copy())
        encoded_quantized = quantization(encoded, 8)
        fig = plt.figure(figsize=(8, 10))
        plt.plot(x_range, encoded_quantized, label='Sygnal po kompresji a-law po kwantyzacji do 8-bitów')
        plt.plot(x_range, encoded, label='Sygnal po kompresji a-law bez kwantyzacji')

        encoded = mu_law_encode(x_range.copy())
        ecoded_quantized = quantization(encoded, 8)
        plt.plot(x_range, ecoded_quantized, label='Sygnal po kompresji mu-law po kwantyzacji do 8-bitów')
        plt.plot(x_range, encoded, label='Sygnal po kompresji mu-law bez kwantyzacji')
        plt.title("Krzywa kompresji")
        plt.xlabel("Wartość syngału wejściowego")
        plt.ylabel("Wartość sygnału wyjściowego")
        plt.legend(loc='upper left')

        fig.tight_layout()
        memfile = BytesIO()
        fig.savefig(memfile)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close(fig)

        fig = plt.figure(figsize=(8, 10))
        plt.plot(x_range, x_range, label='Sygnał oryginalny')
        plt.plot(x_range, A_law_decode(encoded_quantized.copy()), label='Sygnał po dekompresji z a-law (kwantyzacja 8-bitów)')
        plt.plot(x_range, mu_law_decode(ecoded_quantized.copy()), label='Sygnał po dekompresji z mu-law (kwantyzacja 8-bitów)')
        plt.plot(x_range, quantization(x_range.copy(), 8), label='Sygnał oryginalny po kwantyzacji do 8 bitów')
        plt.title("Krzywa po dekompresji")
        plt.xlabel("Wartość syngału wejściowego")
        plt.ylabel("Wartość sygnału wyjściowego")
        plt.legend(loc='upper left')

        fig.tight_layout()
        memfile = BytesIO()
        fig.savefig(memfile)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close(fig)

        document.add_page_break()

    for x_range in [np.linspace(-1, 1, 1000), np.linspace(-0.5, -0.25, 1000)]:
        signal = 0.9 * np.sin(np.pi * x_range * 4)
        fig = plt.figure(figsize=(8, 7))
        plt.suptitle("Przykład A kwantyzacja do 6 bitów")

        plt.subplot(5, 1, 1)
        plt.title("Sygnal oryginalny")
        plt.plot(x_range, signal)

        plt.subplot(5, 1, 2)
        plt.title("Kompresja A-law")
        plt.plot(x_range, A_law_decode(quantization(A_law_encode(signal.copy()), 6)), label='Kompresja A-law')

        plt.subplot(5, 1, 3)
        plt.title("Kompresja mu-law")
        plt.plot(x_range, mu_law_decode(quantization(mu_law_encode(signal.copy()), 6)), label='Kompresja mu-law')

        plt.subplot(5, 1, 4)
        plt.title("Kompresja DPCM bez predykcji")
        plt.plot(x_range, DPCM_decode(DPCM_encode(signal, 6)), label='Kompresja DPCM bez predykcji')

        plt.subplot(5, 1, 5)
        plt.title("Kompresja DPCM z predykcją")
        plt.plot(x_range, DPCM_decode_prediction(DPCM_encode_prediction(signal.copy(), 6, n=3), n=3), label='Kompresja DPCM z predykcją')

        plt.tight_layout()
        memfile = BytesIO()
        fig.savefig(memfile)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close(fig)

        document.add_page_break()

        fig = plt.figure(figsize=(12, 6))
        plt.title('Przykład B kwantyzacja do 6 bitów')
        plt.plot(x_range, signal, label='Sygnał oryginalny')
        plt.plot(x_range, A_law_decode(quantization(A_law_encode(signal.copy()), 6)), label='Sygnał po dekompresji z a-law')
        plt.plot(x_range, mu_law_decode(quantization(mu_law_encode(signal.copy()), 6)), label='Sygnał po dekompresji z mu-law')
        plt.plot(x_range, DPCM_decode(DPCM_encode(signal, 6)), label='Sygnał po dekompresji z DPCM bez predykcji')
        plt.plot(x_range, DPCM_decode_prediction(DPCM_encode_prediction(signal.copy(), 6, n=3), n=3), label='Sygnał po dekompresji z DPCM z predykcją')
        plt.legend(loc='upper left')

        plt.tight_layout()
        memfile = BytesIO()
        fig.savefig(memfile)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close(fig)



    docx_path = 'report.docx'
    document.save(docx_path)







# a()
# a_2()
# b()
generate_report()
