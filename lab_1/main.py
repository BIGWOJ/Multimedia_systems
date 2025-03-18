import numpy as np
import matplotlib.pyplot as plt
import sounddevice as sd
import soundfile as sf
import scipy.fftpack

def task_1():
    data, fs = sf.read('SOUND_INTRO/sound1.wav', dtype='float32')
    # sd.play(data, fs)
    # status = sd.wait()

    sound_L = data[:, 0]
    sound_R = data[:, 1]
    sound_mix = (sound_L + sound_R) / 2

    # sf.write('SOUND_INTRO/sound1_L.wav', sound_L, fs)
    # sf.write('SOUND_INTRO/sound1_R.wav', sound_R, fs)
    # sf.write('SOUND_INTRO/sound1_mix.wav', sound_mix, fs)
    #
    # x = np.arange(0, data.shape[0]) / fs
    #
    # plt.subplot(2,1,1)
    # plt.plot(x, data[:,0])
    #
    # plt.subplot(2,1,2)
    # plt.plot(x, data[:,1])
    # plt.show()
    #

    data, fs = sf.read('SIN/sin_440Hz.wav', dtype=np.int32)
    plt.figure()
    plt.subplot(2, 1, 1)
    plt.plot(np.arange(0, data.shape[0]) / fs, data)

    plt.subplot(2, 1, 2)
    yf = scipy.fftpack.fft(data)
    plt.plot(np.arange(0, fs, 1.0 * fs / (yf.size)), np.abs(yf))
    plt.show()

    # Decibel scale
    fsize = 2**8
    plt.figure()
    plt.subplot(2, 1, 1)
    plt.plot(np.arange(0, data.shape[0]) / fs, data)
    plt.subplot(2, 1, 2)
    yf = scipy.fftpack.fft(data, fsize)
    plt.plot(np.arange(0, fs / 2, fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])))
    plt.show()

def task_2():
    from useful_modules import generate_figures
    signal, fs = sf.read('SIN/sin_440Hz.wav', dtype='float32')
    time_margin = [0.0, 0.005]
    generate_figures(signal, fs, time_margin)

def task_3():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO

    files = ['SIN/sin_60Hz.wav', 'SIN/sin_440Hz.wav', 'SIN/sin_8000Hz.wav']
    document = Document()
    document.add_heading('Change title', 0)

    files = ['SIN/sin_60Hz.wav', 'SIN/sin_440Hz.wav', 'SIN/sin_8000Hz.wav']
    margins = [[0, 0.02], [0.133, 0.155]]
    fsizes = [2**8, 2**12, 2**16]

    for file in files:
        document.add_heading('Plik - {}'.format(file), 2)
        signal, fs = sf.read(file, dtype='float32')
        for fsize in fsizes:
            document.add_heading('Fsize: {}'.format(fsize), 3)
            fig, axs = plt.subplots(2, 1, figsize=(10, 7))

            # Analizing code, plots
            ############################################################

            signal_part = signal[:fsize]
            fft_values = np.fft.fft(signal_part)
            freqs = np.fft.fftfreq(len(signal_part), d=1/fs)
            magnitudes = np.abs(fft_values)

            peak_index = np.argmax(magnitudes[:len(magnitudes)//2])
            peak_freq = freqs[peak_index]
            peak_amplitude = magnitudes[peak_index]

            axs[0].plot(signal_part)
            axs[0].set_title("Sygnał")
            axs[0].set_xlabel("Próbka")

            axs[1].plot(freqs[:len(freqs) // 2], magnitudes[:len(magnitudes) // 2])
            axs[1].set_title("Widmo")
            axs[1].set_xlabel("Hz")
            axs[1].set_ylabel("Amplituda")

            ############################################################

            fig.tight_layout(pad=1.5)
            memfile = BytesIO()
            fig.savefig(memfile)

            document.add_picture(memfile, width=Inches(6))

            memfile.close()

            # Text data - values, outputs etc
            ############################################################
            document.add_paragraph('Najwyższa częstotliwość = {:.0f}'.format(peak_freq))
            document.add_paragraph('Wartość amplitudy dla {:.0f}Hz = {:.0f}'.format(peak_freq, peak_amplitude))
            ############################################################

    document.save('report.docx')

# task_1()(
# task_2()
task_3()