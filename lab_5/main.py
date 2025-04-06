import numpy as np
from useful_modules import generate_figures
import soundfile as sf
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def quantize(data, bit, m=None, n=None):
    dtype = data.dtype

    if not np.issubdtype(dtype, np.floating):
        DataF = data.astype(float)
    else:
        DataF = data.copy()

    if m is None or n is None:
        if np.issubdtype(dtype, np.integer):
            info = np.iinfo(dtype)
            m = info.min
            n = info.max
        else:
            m = DataF.min()
            n = DataF.max()

    d = 2 ** bit - 1
    DataF = (DataF - m) / (n - m)
    DataF = np.round(DataF * d) / d
    DataF = DataF * (n - m) + m

    return DataF.astype(data.dtype)

def decimation(data, factor, fs):
    new_fs = fs / factor
    return data[::factor], new_fs

def interpolation(data, fs, fs_new, method='linear'):
    from scipy.interpolate import interp1d

    n = len(data)
    duration = n / fs
    n_new = int(duration * fs_new)

    t = np.linspace(0, duration, n)
    t_new = np.linspace(0, duration, n_new)

    if method == 'linear':
        interpolated = interp1d(t, data)
    elif method == 'cubic':
        interpolated = interp1d(t, data, kind='cubic')

    data_new = interpolated(t_new)

    return data_new.astype(data.dtype)

def test_quantize(signal, fs):
    generate_figures(signal, fs)
    generate_figures(quantize(signal, 1), fs)
    generate_figures(quantize(signal, 2), fs)
    generate_figures(quantize(signal, 4), fs)
    generate_figures(quantize(signal, 8), fs)

def test_decimation(signal, fs):
    # generate_figures(signal, fs)
    # decimated_signal, decimated_fs = decimation(signal, 2, fs)
    # generate_figures(decimated_signal, decimated_fs)
    # decimated_signal, decimated_fs = decimation(signal, 4, fs)
    # generate_figures(decimated_signal, decimated_fs)
    # decimated_signal, decimated_fs = decimation(signal, 5, fs)
    # generate_figures(decimated_signal, decimated_fs, time_margin=[0, 0.002])
    # decimated_signal, decimated_fs = decimation(signal, 10, fs)
    # generate_figures(decimated_signal, decimated_fs, [0, 0.002])
    decimated_signal, decimated_fs = decimation(signal, 2, fs)
    generate_figures(decimated_signal, decimated_fs)

def test_interpolation(signal, fs):
    fs_new = 2000
    signal_lin = interpolation(signal, fs, fs_new, method='linear')
    signal_cubic = interpolation(signal, fs, fs_new, method='cubic')
    # generate_figures(signal, fs)
    generate_figures(signal_lin, fs_new, [0, 1])
    generate_figures(signal_cubic, fs_new, [0, 1])

def test_sing_files(document=None, test_sing=False):
    import sounddevice as sd

    files = ['SING/sing_low1.wav', 'SING/sing_medium1.wav', 'SING/sing_high1.wav']
    bits = [4, 8]
    decimation_steps = [4, 6, 10, 24]
    interpolation_fs = [4000, 8000, 11999, 16000, 16953]
    interpolation_methods = ['linear', 'cubic']

    if test_sing:
        for file in files:
            print(f"\n===== Original sound of {file} =====")
            signal, fs = sf.read(file)
            sd.play(signal, fs)
            status = sd.wait()

            for bit in bits:
                print(f"\t=== Quantization {bit} bits ===")
                quantized_signal = quantize(signal, bit)
                sd.play(quantized_signal, fs)
                status = sd.wait()

            for step in decimation_steps:
                print(f"\t=== Decimation {step} steps ===")
                decimated_signal, decimated_fs = decimation(signal, step, fs)
                sd.play(decimated_signal, decimated_fs)
                status = sd.wait()

            for fs_new in interpolation_fs:
                print(f"\t=== Interpolation to {fs_new} Hz ===")
                for method in interpolation_methods:
                    print(f"\t\t=== Method {method} ===")
                    interpolated_signal = interpolation(signal, fs, fs_new, method)
                    sd.play(interpolated_signal, fs_new)
                    status = sd.wait()

    if document:
        table = document.add_table(rows=4, cols=4)

        text_quant_low = ('Kwantyzacja do 4 bitów - dźwięk nie był słyszalny.\n'
                          'Kwantyzacja do 8 bitów - dźwięk był znacząco bardziej szumiący względem oryginału.\n')

        text_quant_med = ('Kwantyzacja do 4 bitów - dźwięk "ostry", znacząco głośniejszy, nieprzyjmeny dla ucha.\n'
                          'Kwantyzacja do 8 bitów - dźwięk bliższy oryginalnemu, pojawiły się szumy.\n')

        text_quant_high = ('Kwantyzacja do 4 bitów - dźwięk "ostry", znacząco głośniejszy, nieprzyjmeny dla ucha, pojawienie się dziwnych zakłóceń w tle.\n'
                           'Kwantyzacja do 8 bitów - dźwięk bliższy oryginalnemu, pojawiły się szumy.\n')

        text_dec_low = ('Decymacje z krokiem 4 oraz 6 brzmią bardzo podobnie - wyraźny dźwięk, czysty, bliski oryginalnemu.\n'
                        'Decymacja z krokiem 10  bardzo zniekształca dźwięk, mocno go zniżając, tłumi go.\n'
                        'Decymacja z krokiem 24 najmocniej tłumi dźwięk, jest on bardzo niski,  prawie nie jest możliwy do zrozumienia, dodatkowo pojawiają się lekkie szumy w tle.\n')

        text_dec_med = ('Decymacje z krokiem 4 oraz 6 brzmią bardzo podobnie - wyraźny dźwięk, czysty, bliski oryginalnemu.\n'
                        'Decymacja z krokiem 10 znacząco zniekształca dźwięk, pojawiają się dziwne, nieoryginalne dźwięki, zakłócenia.\n'
                        'Decymacja z krokiem 24 najbardziej zniekształca dźwięk, względem kroku 10 pojawiają się dodatkowe niepożądane dźwięki.')

        text_dec_high = 'Interpretacja niemalże identyczna jak w przypadku medium1 z tą różnicą, że dodatkowe niepożądane dźwięki pojawiają się już przy kroku 6.'

        text_inter_low = ('Nie usłyszałem żadnych różnic w dźwięku, dźwięk był czysty, bliski oryginalnemu we wszystkich przypadkach oprócz interpolacji do 4000Hz.\n'
                          'Interpolacja do 4000Hz - dźwięk był bardzo zniekształcony, obniżony, stłumiony, pojawiły się drobne szumy.\n'
                          'Dodatkowo nie zaobserwowałem żadnych różnic w dźwięku pomiędzy interpolacją liniową a sześcienną (nieliniową).\n')

        text_inter_med = ('Interpolacja do 4000Hz - dźwięk znacząco zniekształcony, bardzo dużo niepożądanych dźwięków, nieprzyjemny dla ucha, drastycznie różni się od oryginalnego.\n'
                          'Interpolacja do 8000Hz - podobnie jak w przypadku 4000Hz, natomiast skala zniekształcenia jest znacząco mniejsza,\n'
                          'dźwięk znacząco odbiega od oryginalnego brzmienia.\n'
                          'Interpolacje do 11999Hz - dźwięk bardzo podobny do oryginalnego, jednakże wyczuwana jest minimalna różnica.\n'
                          'Interpolacje do 16000Hz oraz 16953Hz - dźwięk w moim odczuciu niemalże identyczny do oryginalnego, nie zanotowałem żadnych różnic.\n'
                          'Podcza interpolacji sześciennej (nieliniowej) odczułem minimalną poprawę odwzorowania względem liniowej.')

        text_inter_high = ('Interpolacja do 4000Hz - dźwięk znacząco zniekształcony, pojawiają się niepożądane dźwięki, dziwny szum przez cały czas.\n'
                           'Interpolacja do 8000Hz - dźwięk mniej zniekształcony względem 4000Hz, mniej niepożądanych dźwięków, jednakże wciąż odbiegający od oryginalnego.\n'
                           'Interpolacje do 11999Hz, 16000Hz oraz 16953Hz - dźwięk w moim odczuciu niemalże identyczny do oryginalnego, nie słyszałem żadnych różnic.\n'
                           'Podczas interpolacji szceściennej (nieliniowej) odczułem minimalną poprawę odwzorowania względem liniowej.')

        cell_texts = [
            '', 'Kwantyzacja', 'Decymacja', 'Interpolacja',
            'Plik low1', text_quant_low, text_dec_low, text_inter_low,
            'Plik medium1', text_quant_med, text_dec_med, text_inter_med,
            'Plik high1', text_quant_high, text_dec_high, text_inter_high
        ]

        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                cell = table.cell(row, col)
                cell.text = cell_texts[row * len(table.columns) + col]
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '5')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

                if col == 0:
                    tcPr.append(OxmlElement('w:tcW'))
                    tcPr[-1].set(qn('w:w'), '1000')
                    tcPr[-1].set(qn('w:type'), 'dxa')
                else:
                    tcPr.append(OxmlElement('w:tcW'))
                    tcPr[-1].set(qn('w:w'), '5000')
                    tcPr[-1].set(qn('w:type'), 'dxa')

        return document

def generate_report(test_sing=False):
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert
    from os import remove
    import matplotlib.pyplot as plt

    files = ['SIN/sin_60Hz.wav', 'SIN/sin_440Hz.wav', 'SIN/sin_8000Hz.wav', 'SIN/sin_combined.wav']
    bits = [4, 8, 16, 24]
    decimation_steps = [2, 4, 6, 10, 24]
    interpolation_fs = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
    interpolation_methods = ['linear', 'cubic']
    time_margins = [[0, 0.05], [0, 0.006], [0, 0.0004], [0, 0.0085]]
    time_margins_interpolation_8000 = [[0, 1], [0, 1], [0, 1], [0, 0.001], [0, 0.01], [0, 0.005], [0, 0.0005], [0, 0.001]]

    document = Document()
    document.add_heading('Kwantyzacja i próbkowanie dźwięku oraz re-sampling\nWojciech Latos', 0)

    for time_margin_index, file in enumerate(files):
        print(file)
        time_margin = time_margins[time_margin_index]
        document.add_heading('Sygnał - {}'.format(file), 2)
        signal, fs = sf.read(file)

        document.add_heading('Sygnał oryginalny', 3)
        fig = generate_figures(signal, fs, time_margin, True)
        fig_memfile = BytesIO()
        fig.savefig(fig_memfile)
        fig_memfile.seek(0)
        document.add_picture(fig_memfile, width=Inches(7))
        fig_memfile.close()
        plt.close(fig)

        for bit in bits:
            document.add_heading('Kwantyzacja {} bit'.format(bit), 3)
            fig = generate_figures(quantize(signal, bit), fs, time_margin, True)
            fig_memfile = BytesIO()
            fig.savefig(fig_memfile)
            fig_memfile.seek(0)

            document.add_picture(fig_memfile, width=Inches(7))

            fig_memfile.close()
            plt.close(fig)

        for step in decimation_steps:
            if file.__contains__('8000'):
                if step in [4, 10]:
                    time_margin = np.array(time_margin) * 3
            document.add_heading('Decymacja {} kroki'.format(step), 3)
            decimated_signal, decimated_fs = decimation(signal, step, fs)
            fig = generate_figures(decimated_signal, decimated_fs, time_margin, True)
            fig_memfile = BytesIO()
            fig.savefig(fig_memfile)
            fig_memfile.seek(0)

            document.add_picture(fig_memfile, width=Inches(7))

            fig_memfile.close()
            plt.close(fig)

        for fs_new in interpolation_fs:
            document.add_heading('Interpolacja do {}Hz'.format(fs_new), 3)
            for method in interpolation_methods:
                if file.__contains__('8000'):
                    time_margin = time_margins_interpolation_8000[interpolation_fs.index(fs_new)]
                document.add_heading('Metoda {}'.format(method), 4)
                interpolated_signal = interpolation(signal, fs, fs_new, method)
                fig = generate_figures(interpolated_signal, fs_new, time_margin, generate_report=True)
                fig_memfile = BytesIO()
                fig.savefig(fig_memfile)
                fig_memfile.seek(0)

                document.add_picture(fig_memfile, width=Inches(7))

                fig_memfile.close()
                plt.close(fig)

    document.add_page_break()
    document.add_heading('Wnioski z testów plików SIN', 1)
    table = document.add_table(rows=2, cols=4)

    text_sin_60 = ('Im wyższa liczba bitów podczas kwantyzacji, tym wykres amplitudowy jest bardziej zbliżony oryginalnemu, '
                   'wartości decybeli na widmie zbiegają do tych z wykresu niezmodyfikowanego sygnału.\n'
                   'Podczas decymacji z różnymi krokami ogólny kształt sygnału został zachowany.\n'
                   'Im wyższa wartość interpolacji tym większe/mniejsze są maksymalne/minimalnie wartości decybeli na skali.\n'
                   'Dodatkowo początek wykresu jest bardziej strzelisty, metody interpolacji nie różnią się od siebie.\n')

    text_sin_440 = ('Im wyższa liczba bitów podczas kwantyzacji, tym wykres amplitudowy jest bardziej zbliżony oryginalnemu, '
                   'wartości decybeli na widmie zbiegają do tych z wykresu niezmodyfikowanego sygnału.\n'
                    'Podczas decymacji ogólny kształt pozostał taki sam, jednakże wykres stał się bardziej "kwadratowy", widać wyraźne linie, nie jest gładki.\n'
                    'Im wyższa wartość interpolacji tym wykres stawał się bardziej gładki, coraz bardziej zbliżony do oryginalnego.\n'
                    'Czubek wykresu skali decybelowej był w tym samym miejscu przez wszystkie parametry interpolacji, natomiast zmieniała się minimalna/maksymalna wartość wykresu.\n'
                    'Przy największych wartościach interpolacji użycie metody sześciennej (nieliniowiej) skutkowało mniej poszarpanym, gładszym wykresem w skali decybelowej.\n')

    text_sin_8000 = ('Wszystkie parametry kwantyzacji skutkowały takim samym wykresem decybelowym.\n'
                     'Podczas decymacji z krokami 10 i 24 wykresy amplitudowe oraz decybelowe są linią prostą, ponieważ zostały wybrane akurat zerowe wartości sygnału.\n'
                     'Największe różnice były podczas interpolacji, gdzie przy metodzie sześciennej (nieliniowej) wykres był sinusoidalny, przy liniowej - trapezowaty.\n'
                     'Wykres decybelowy w przypadku metody sześciennej także jest bardziej gładki, nie jest poszarpany jak ten z metody liniowej.')

    text_sin_combined = ('Przez wszystkie parametry kwantyzacji wykres amplitudowy jak i decybelowy był niezmienny.\n'
                         'Im wyższa wartość decymacji tym o wiele łatwiej o złą analizę sygnału, wykres amplitudowy jest pozbawiony coraz większej ilości szczegółów, wykres jest uogólniany co może mieć zły wpływ na interpretację.\n'
                         'Im wyższa wartość interpolacji tym wykres decybelowy staje się "gęstszy", przekazuje więcej informacji, wykres amplitudowy nabiera coraz więcej szczegółów.')

    cell_texts = [
        'Sin 60Hz ', 'Sin 440Hz', 'Sin 8000Hz', 'Sin combined',
        text_sin_60, text_sin_440, text_sin_8000, text_sin_combined,
    ]

    for row in range(len(table.rows)):
        for col in range(len(table.columns)):
            cell = table.cell(row, col)
            cell.text = cell_texts[row * len(table.columns) + col]
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '5')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    docx_path = 'report.docx'

    if test_sing:
        document.add_page_break()
        document.add_heading('Wnioski z testów plików SING', 1)
        document_with_table = test_sing_files(document)
        document_with_table.save(docx_path)

    else:
        document.save(docx_path)

    convert(docx_path, 'report.pdf')
    remove(docx_path)


# signal, fs = sf.read('SIN/sin_8000Hz.wav')
# test_quantize(signal, fs)
# test_decimation(signal, fs)
# test_interpolation(signal, fs)
# decimated_signal, decimated_fs = decimation(signal, 24, fs)
# generate_figures(decimated_signal, decimated_fs)
# test_sing_files()
generate_report(True)

