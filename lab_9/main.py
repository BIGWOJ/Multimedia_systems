import cv2
import numpy as np
import matplotlib.pyplot as plt
import os

class data:
    def init(self):
        self.Y = None
        self.Cb = None
        self.Cr = None


def Chroma_subsampling(L, subsampling):
    if subsampling == "4:4:4":
        return L
    elif subsampling == "4:2:2":
        return L[:, ::2]
    elif subsampling == "4:4:0":
        return L[::2, :]
    elif subsampling == "4:2:0":
        return L[::2, ::2]
    elif subsampling == "4:1:1":
        return L[:, ::4]
    elif subsampling == "4:1:0":
        return L[::4, :]


def Chroma_resampling(L, subsampling):
    if subsampling == "4:4:4":
        return L
    elif subsampling == "4:2:2":
        return np.repeat(L, 2, axis=1)
    elif subsampling == "4:4:0":
        return np.repeat(L, 2, axis=0)
    elif subsampling == "4:2:0":
        return np.repeat(np.repeat(L, 2, axis=0), 2, axis=1)
    elif subsampling == "4:1:1":
        return np.repeat(L, 4, axis=1)
    elif subsampling == "4:1:0":
        return np.repeat(L, 4, axis=0)


def frame_image_to_class(frame, subsampling):
    Frame_class = data()
    Frame_class.Y = frame[:, :, 0].astype(int)
    Frame_class.Cb = Chroma_subsampling(frame[:, :, 2].astype(int), subsampling)
    Frame_class.Cr = Chroma_subsampling(frame[:, :, 1].astype(int), subsampling)
    return Frame_class


def frame_layers_to_image(Y, Cr, Cb, subsampling):
    Cb = Chroma_resampling(Cb, subsampling)
    Cr = Chroma_resampling(Cr, subsampling)
    return np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8)


def compress_KeyFrame(Frame_class):
    KeyFrame = data()
    KeyFrame.Y = Frame_class.Y
    KeyFrame.Cb = Frame_class.Cb
    KeyFrame.Cr = Frame_class.Cr
    return KeyFrame


def decompress_KeyFrame(KeyFrame, subsampling):
    Y = KeyFrame.Y
    Cb = KeyFrame.Cb
    Cr = KeyFrame.Cr
    frame_image = frame_layers_to_image(Y, Cr, Cb, subsampling)
    return frame_image


def compress_not_KeyFrame(Frame_class, KeyFrame, divisor):
    Compress_data = data()
    Compress_data.Y = ((Frame_class.Y - KeyFrame.Y) // divisor).astype(np.int16)
    Compress_data.Cb = ((Frame_class.Cb - KeyFrame.Cb) // divisor).astype(np.int16)
    Compress_data.Cr = ((Frame_class.Cr - KeyFrame.Cr) // divisor).astype(np.int16)
    return Compress_data

def decompress_not_KeyFrame(Compress_data, KeyFrame, divisor, subsampling):
    Y = KeyFrame.Y + (Compress_data.Y * divisor)
    Cb = KeyFrame.Cb + (Compress_data.Cb * divisor)
    Cr = KeyFrame.Cr + (Compress_data.Cr * divisor)
    return frame_layers_to_image(Y, Cr, Cb, subsampling)

def run_length_encode(data):
    flat = data.flatten()
    values = []
    counts = []
    prev = flat[0]
    count = 1
    for i in range(1, len(flat)):
        if flat[i] == prev:
            count += 1
        else:
            values.append(prev)
            counts.append(count)
            prev = flat[i]
            count = 1
    values.append(prev)
    counts.append(count)
    return (np.array(values), np.array(counts))


def plotDiffrence(frame_original, frame_decompressed, roi, subsampling, divisor, KeyFrame_counter, frame_counter):
    x, y, w, h = roi
    roi_original = frame_original[y:y + h, x:x + w, :]
    roi_decompressed = frame_decompressed[y:y + h, x:x + w, :]
    diff = roi_decompressed.astype(float) - roi_original.astype(float)

    fig, axs = plt.subplots(4, 3, figsize=(10, 8))
    fig.suptitle("Subsampling: {}, Divisor: {}, KeyFrame counter: {}, Frame: {}".format(subsampling, divisor, KeyFrame_counter, frame_counter), fontsize=16)

    axs[0][0].imshow(roi_original)
    axs[0][0].set_title('Original (ROI)')


    axs[0][1].imshow(roi_decompressed)
    axs[0][1].set_title('Decompressed (ROI)')
    axs[1][0].imshow(roi_original[:, :, 0], cmap='Reds')
    axs[1][0].set_title('Original R')
    axs[1][1].imshow(roi_decompressed[:, :, 0], cmap='Reds')
    axs[1][1].set_title('Decompressed R')
    axs[1][2].imshow(diff[:, :, 0], cmap='Reds', vmin=-255, vmax=255)
    axs[1][2].set_title('Delta R')

    axs[2][0].imshow(roi_original[:, :, 1], cmap='Greens')
    axs[2][0].set_title('Original G')
    axs[2][1].imshow(roi_decompressed[:, :, 1], cmap='Greens')
    axs[2][1].set_title('Decompressed G')
    axs[2][2].imshow(diff[:, :, 1], cmap='Greens', vmin=-255, vmax=255)
    axs[2][2].set_title('Delta G')

    axs[3][0].imshow(roi_original[:, :, 2], cmap='Blues')
    axs[3][0].set_title('Original B')
    axs[3][1].imshow(roi_decompressed[:, :, 2], cmap='Blues')
    axs[3][1].set_title('Decompressed B')
    axs[3][2].imshow(diff[:, :, 2], cmap='Blues', vmin=-255, vmax=255)
    axs[3][2].set_title('Delta B')

    for ax in axs.flat:
        ax.axis('off')

    plt.tight_layout()
    return fig

def generate_report():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert
    from os import remove
    catalog = './video'
    clip = "clip_4.mp4"
    frames_count = 15 #<0 - entire video, >0 - number of frames
    key_frame_counter = 3  # every key_frame_counter frame is a key frame
    plot_frames = np.array([2, 11])  # automatically draw plots
    auto_pause_frames = np.array([25])  # automatically pause
    show_frames = False
    ROI = [[200, 300, 150, 250]]

    cap = cv2.VideoCapture(os.path.join(catalog, clip))

    if frames_count < 0:
        frames_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

    cv2.namedWindow('Normal Frame')
    cv2.namedWindow('Decompressed Frame')

    compression_information = np.zeros((3, frames_count))

    document = Document()
    document.add_heading('Kwantyzacja i próbkowanie dźwięku oraz re-sampling\nWojciech Latos', 0)
    document.add_heading('Podsumowanie cz. 1', level=1)
    document.add_paragraph('Na podstawie poniższych wykresów można zauważyć, że najbardziej obiecującymi parametrami, które moim zdaniem mają największe szanse na dalszą lepszą kompresję w drugiej części zadania'
                           'są :\nSubsampling (ex aequo): 4:2:0, 4:1:1 oraz 4:1:0.\n'
                           'Sądzę, że 4:1:0 jest najbardziej obiecujacym wyborem pod względem przyszłej kolejnej kompresji, ponieważ jest to najmocniejsza kompresja chrominancji. Z zasady działania subsamplingu sądzę, że będzie to najlepszy wybór.\n'
                           'Dzielnik oraz licznik klatek kluczowych w owym teście nie wskazywał żadnych wpływów na jakość kompresji.\n'
                           'Gdybyśmy potrzebowali tak zwanej pixel-perfect jakości wideo, w dużym stopniu nie zwracając uwagi na rozmiar skompresowanego pliku, wybrałym  opcje 4:2:2 bądź 4:4:0.\n'
                           'Opcja 4:4:4 nie dokonuje kompresji. Wykres jest linią prostą o zerowej wartości.')
    document.add_page_break()

    document.add_heading('Podsumowanie cz. 2', level=1)
    document.add_paragraph('Mając do dyspozycji tylko poniższe wykresy z drugiej części badania wysnuwam wniosek, że najlepszą odległością pomiędzy klatkami kluczowymi okazała się liczba 2.\n'
                           'W przypadku warstw G oraz B stopień kompresji oscylował w okolicach wartości 90-95%, natomiast w przypadku warstwy R wartość ta wynosiła prawie 65% do 85%.\n'
                           'Spadki kompresji na wykresach zbiegają się z momentami, w których została ustalona nowa klatka kluczowa.\n'
                           'Mając jako kryterium tylko rozmiar skompresowanego pliku, moglibyśmy w celach wyłonienia najlepszej metody obliczyć pole pod krzywą. Im większe pole, tym kompresja okazałaby się lepszą.\n'
                           'Bazując na poniższych wykresach oraz przedstawionej metodzie oceny można byłoby wybrać odległość między klatkami kluczowymi na poziomie 12 klatek.\n'
                           'Jednakże teraz wchodzi sprawa jakości samego obrazu po dekompresji. W przypadku 12 klatek kluczowych jakość wideo jest znacznie gorsza niż w przypadku 2 klatek kluczowych.\n'
                           'Uważam, że optymalnym rozwiązaniem jest wybranie odległości co 3 klatki dla wideo o długości 15 klatek.\n'
                           'Dodatkowo można zauważyć, że kanały G i B (zielony i niebieski), które są bezpośrednio związane z luminancją w modelu YCbCr, osiągają wysokie wartości kompresji na poziomie około 90-95%. To sugeruje, że te dane dobrze się je kompresuje strumieniowo')
    document.add_page_break()

    document.add_heading(f'Badanie cz. 1 - bez RLE', level=1)
    for subsampling in ['4:4:4', '4:2:2', '4:4:0', '4:2:0',  '4:1:1', '4:1:0']:
        for divisor in [1, 2, 4, 8]:
            for i in range(frames_count):
                ret, frame = cap.read()
                if show_frames:
                    cv2.imshow('Normal Frame', frame)
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                Frame_class = frame_image_to_class(frame, subsampling)
                if (i % key_frame_counter) == 0:
                    KeyFrame = compress_KeyFrame(Frame_class)
                    cR = KeyFrame.Y
                    cG = KeyFrame.Cb
                    cB = KeyFrame.Cr
                    Decompresed_Frame = decompress_KeyFrame(KeyFrame, subsampling)
                else:
                    Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame, divisor)
                    cR = Compress_data.Y
                    cG = Compress_data.Cb
                    cB = Compress_data.Cr
                    Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame, divisor, subsampling)

                compression_information[0,i]= (frame[:,:,0].size - cR.size)/frame[:,:,0].size
                compression_information[1,i]= (frame[:,:,0].size - cG.size)/frame[:,:,0].size
                compression_information[2,i]= (frame[:,:,0].size - cB.size)/frame[:,:,0].size

                if show_frames:
                    cv2.imshow('Decompressed Frame',cv2.cvtColor(Decompresed_Frame,cv2.COLOR_YCrCb2BGR))

                if np.any(plot_frames == i):  # rysuj wykresy
                    for r in ROI:
                        fig = plotDiffrence(frame, Decompresed_Frame, r, subsampling, divisor, key_frame_counter, i)
                        fig_memfile = BytesIO()
                        fig.savefig(fig_memfile)
                        fig_memfile.seek(0)
                        document.add_picture(fig_memfile, width=Inches(7))
                        fig_memfile.close()
                        plt.close(fig)

                if np.any(auto_pause_frames == i):
                    cv2.waitKey(-1)  # wait until any key is pressed

                k = cv2.waitKey(1) & 0xff

                if k == ord('q'):
                    break
                elif k == ord('p'):
                    cv2.waitKey(-1)  # wait until any key is pressed

            fig = plt.figure()
            plt.plot(np.arange(0, frames_count), compression_information[0, :] * 100, label='R')
            plt.plot(np.arange(0, frames_count), compression_information[1, :] * 100, label='G')
            plt.plot(np.arange(0, frames_count), compression_information[2, :] * 100, label='B')
            plt.title("{}, subsampling: {}, divisor: {}, KeyFrame counter:{}".format(clip, subsampling, divisor, key_frame_counter))
            plt.xlabel('Frame number')
            plt.ylabel('Compression ratio [%]')
            plt.legend()

            fig.tight_layout()
            memfile = BytesIO()
            fig.savefig(memfile)
            memfile.seek(0)
            document.add_picture(memfile, width=Inches(6))
            memfile.close()
            plt.close(fig)


    document.add_page_break()
    document.add_heading(f'Badanie cz. 2 - z RLE', level=1)
    subsampling = '4:1:0'
    divisor = 8

    for key_frame_counter in [2, 3, 5, 8, 12]:
        for i in range(frames_count):
            ret, frame = cap.read()
            if show_frames:
                cv2.imshow('Normal Frame', frame)
            frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            Frame_class = frame_image_to_class(frame, subsampling)
            if (i % key_frame_counter) == 0:
                KeyFrame = compress_KeyFrame(Frame_class)
                cR = KeyFrame.Y
                cG = KeyFrame.Cb
                cB = KeyFrame.Cr
                Decompresed_Frame = decompress_KeyFrame(KeyFrame, subsampling)
            else:
                Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame, divisor)
                cR = Compress_data.Y
                cG = Compress_data.Cb
                cB = Compress_data.Cr
                Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame, divisor, subsampling)

            if show_frames:
                cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))

            compression_information[0, i] = (frame[:, :, 0].size - len(run_length_encode(cR)[0])) / frame[:, :, 0].size
            compression_information[1, i] = (frame[:, :, 0].size - len(run_length_encode(cG)[0])) / frame[:, :, 0].size
            compression_information[2, i] = (frame[:, :, 0].size - len(run_length_encode(cB)[0])) / frame[:, :, 0].size

        fig = plt.figure()
        plt.plot(np.arange(0, frames_count), compression_information[0, :] * 100, label='R')
        plt.plot(np.arange(0, frames_count), compression_information[1, :] * 100, label='G')
        plt.plot(np.arange(0, frames_count), compression_information[2, :] * 100, label='B')
        plt.title("{}, subsampling: {}, divisor: {}, KeyFrame counter:{} z RLE".format(clip, subsampling, divisor,
                                                                                       key_frame_counter))
        plt.xlabel('Frame number')
        plt.ylabel('Compression ratio [%]')
        plt.legend()

        fig.tight_layout()
        memfile = BytesIO()
        fig.savefig(memfile)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close(fig)

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')
    remove(docx_path)

generate_report()