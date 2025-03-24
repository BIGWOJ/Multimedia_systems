import matplotlib.pyplot as plt
import numpy as np
import cv2

def scale_nearest_neighbour(img, scale):
    height, width, RGB = img.shape
    new_width = np.ceil(width * scale).astype(int)
    new_height = np.ceil(height * scale).astype(int)

    X = np.linspace(0, width - 1, new_width).astype(int)
    Y = np.linspace(0, height - 1, new_height).astype(int)

    new_img = np.zeros((new_height, new_width, RGB), dtype=np.uint8)

    for i in range(new_height):
        for j in range(new_width):
            new_img[i, j, :] = img[Y[i], X[j], :]

    return new_img

def scale_bilinear(img, scale):
    height, width = img.shape[:2]
    new_height, new_width = int(height * scale), int(width * scale)
    new_img = np.zeros((new_height, new_width, img.shape[2])).astype(img.dtype)

    for i in range(new_height):
        for j in range(new_width):
            x_original = i / scale
            y_original = j / scale

            x0, y0 = int(x_original), int(y_original)
            x1, y1 = np.ceil(x0), np.ceil(y0)

            x_difference, y_difference = x_original - x0, y_original - y0

            Q11 = img[x0, y0]
            Q21 = img[x1, y0]
            Q12 = img[x0, y1]
            Q22 = img[x1, y1]

            new_img[i, j] = (Q11 * (1 - x_difference) * (1 - y_difference) +
                             Q21 * x_difference * (1 - y_difference) +
                             Q12 * (1 - x_difference) * y_difference +
                             Q22 * x_difference * y_difference)

    return new_img

def reduce(img, scale, method="mean"):
    height, width, RGB = img.shape
    new_height, new_width = int(height / scale), int(width / scale)
    new_img = np.zeros((new_height, new_width, img.shape[2])).astype(img.dtype)

    for i in range(new_height):
        for j in range(new_width):
            start_row = int(i * scale)
            end_row = int((i + 1) * scale)
            start_col = int(j * scale)
            end_col = int((j + 1) * scale)
            block = img[start_row:end_row, start_col:end_col]

            if method == "mean":
                new_img[i, j] = np.mean(block, axis=(0, 1))
            elif method == "median":
                new_img[i, j] = np.median(block, axis=(0, 1))
            elif method == "weighted_average":
                block_height, block_width = block.shape[:2]
                weights = np.random.rand(block_height, block_width)
                weights /= np.sum(weights)
                for color in range(RGB):
                    new_img[i, j, color] = np.average(
                        block[:, :, color],
                        weights=weights,
                        axis=(0, 1)
                    )

    return new_img

def get_tests_fragments(img, reduce=False):
    if reduce:
        return [img[:img.shape[0] // 50, :img.shape[1] // 50],
                img[img.shape[0] // 3:img.shape[0] // 3 + img.shape[0] // 50, img.shape[1] // 3:img.shape[1] // 3 + img.shape[1] // 50],
                img[2 * img.shape[0] // 3:2 * img.shape[0] // 3 + img.shape[0] // 48, 2 * img.shape[1] // 3:2 * img.shape[1] // 3 + img.shape[1] // 50]]

    return [img[:img.shape[0] // 10, :img.shape[1] // 10],
            img[img.shape[0] // 3:img.shape[0] // 3 + img.shape[0] // 10, img.shape[1] // 3:img.shape[1] // 3 + img.shape[1] // 10],
            img[2 * img.shape[0] // 3:2 * img.shape[0] // 3 + img.shape[0] // 10, 2 * img.shape[1] // 3:2 * img.shape[1] // 3 + img.shape[1] // 10]]

def test_scale_img(img, scale, method, generate_report=False):
    test_original_fragments = get_tests_fragments(img)

    plt.title("Oryginalny obraz")
    plt.imshow(img)
    plt.tight_layout()

    if not generate_report:
        plt.show()
    plt.close()

    figs = []

    for test_original_fragment in test_original_fragments:
        fig = plt.figure(figsize=(10, 7))
        if method == "nearest":
            scaled_fragment = scale_nearest_neighbour(test_original_fragment, scale)
        elif method == "bilinear":
            scaled_fragment = scale_bilinear(test_original_fragment, scale)

        original_edges = cv2.Canny(test_original_fragment, 100, 200)
        scaled_edges = cv2.Canny(scaled_fragment, 100, 200)

        plt.subplot(2, 2, 1)
        plt.title(f"Fragment oryginalnego obrazu")
        plt.imshow(test_original_fragment)

        plt.subplot(2, 2, 2)
        plt.title(f"Fragment powiększonego obrazu")
        plt.imshow(scaled_fragment)

        plt.subplot(2, 2, 3)
        plt.title(f"Krawędzie oryginalnego obrazu")
        plt.imshow(original_edges, cmap='gray')

        plt.subplot(2, 2, 4)
        plt.title(f"Krawędzie powiększonego obrazu")
        plt.imshow(scaled_edges, cmap='gray')

        plt.tight_layout()

        if generate_report:
            figs.append(fig)

        if not generate_report:
            plt.show()

        plt.close()

    if generate_report:
        return figs

def test_reduce_img(img, scale, method, generate_report=False):
    test_original_fragments = get_tests_fragments(img)

    plt.title("Oryginalny obraz")
    plt.imshow(img)
    plt.tight_layout()

    if not generate_report:
        plt.show()
    plt.close()

    figs = []

    for test_original_fragment in test_original_fragments:
        fig = plt.figure(figsize=(10, 7))
        reduced_fragment = reduce(test_original_fragment, scale, method)

        original_edges = cv2.Canny(test_original_fragment, 100, 200)
        reduced_edges = cv2.Canny(reduced_fragment, 100, 200)

        plt.subplot(2, 2, 1)
        plt.title(f"Fragment oryginalnego obrazu")
        plt.imshow(test_original_fragment)

        plt.subplot(2, 2, 2)
        plt.title(f"Fragment zmniejszonego obrazu")
        plt.imshow(reduced_fragment)

        plt.subplot(2, 2, 3)
        plt.title(f"Krawędzie oryginalnego obrazu")
        plt.imshow(original_edges, cmap='gray')

        plt.subplot(2, 2, 4)
        plt.title(f"Krawędzie powiększonego obrazu")
        plt.imshow(reduced_edges, cmap='gray')

        plt.tight_layout()

        if generate_report:
            figs.append(fig)

        if not generate_report:
            plt.show()

        plt.close()

    if generate_report:
        return figs

def tests():
    img_big = cv2.imread('IMG_BIG/BIG_0004.png')
    img_big = cv2.cvtColor(img_big, cv2.COLOR_BGR2RGB)

    img_small = cv2.imread('IMG_SMALL/SMALL_0001.tif')
    img_small = cv2.cvtColor(img_small, cv2.COLOR_BGR2RGB)
    scale = 1.32

    test_scale_img(img_small, scale, method="nearest")
    test_scale_img(img_small, scale, method="bilinear")
    test_reduce_img(img_big, scale, method="mean")
    test_reduce_img(img_big, scale, method="median")
    test_reduce_img(img_big, scale, method="weighted_average")

def generate_report():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert
    from os import remove

    imgs_small = ['IMG_SMALL/SMALL_0001.tif', 'IMG_SMALL/SMALL_0002.png', 'IMG_SMALL/SMALL_0003.png', 'IMG_SMALL/SMALL_0010.jpg']
    imgs_big = ['IMG_BIG/BIG_0001.jpg', 'IMG_BIG/BIG_0004.png']
    files = imgs_small + imgs_big
    scales_reduces = [1.5, 10, 20] # 50%, 10%, 5%
    scales_upscale = [1.5, 3, 5] # 150%, 300%, 500%

    plt.rcParams.update({'figure.max_open_warning': 0})

    document = Document()
    document.add_heading('Próbkowanie i zmiana rozmiaru obrazu\nWojciech Latos', 0)

    for file in files:
        img = cv2.imread(file)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        document.add_heading('Zdjęcie - {}'.format(file), 1)

        source_memfile = BytesIO()
        plt.imsave(source_memfile, img)
        source_memfile.seek(0)
        document.add_picture(source_memfile, width=Inches(3))
        source_memfile.close()

        if file.__contains__('SMALL'):
            scales = scales_upscale
            operation = "Powiększenie"
            testing_operation= 'scale'
            methods = ['nearest', 'bilinear']
        else:
            scales = scales_reduces
            operation = "Pomniejszenie"
            testing_operation = 'reduce'
            methods = ['mean', 'median', 'weighted_average']

        for scale in scales:
            document.add_heading('{} o x{}'.format(operation, scale), 2)
            for i, method in enumerate(methods):
                document.add_heading('Metoda {}'.format(method), 3)
                if testing_operation == 'scale':
                    generated_figs = test_scale_img(img, scale, method, generate_report=True)

                if testing_operation == 'reduce':
                    generated_figs = test_reduce_img(img, scale, method, generate_report=True)

                for generated_fig in generated_figs:
                    generated_fig_memfile = BytesIO()
                    generated_fig.savefig(generated_fig_memfile, format='png')
                    generated_fig_memfile.seek(0)

                    document.add_picture(generated_fig_memfile, width=Inches(7))

                    generated_fig_memfile.close()
                    plt.close(generated_fig)

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')
    remove(docx_path)

# tests()
generate_report()

