import numpy as np
import cv2


def run_length_encode(img):
    flat = np.array(img).flatten()
    encoded = [img.shape[0], img.shape[1]]
    count = 1

    for i in range(1, len(flat)):
        if flat[i] == flat[i - 1]:
            count += 1
        else:
            encoded.append(count)
            encoded.append(flat[i - 1])
            count = 1
    encoded.append(count)
    encoded.append(flat[-1])

    return np.array(encoded)

def run_length_decode(encoded_data):
    og_shape = encoded_data[:2]
    data = encoded_data[2:]
    decoded = []

    for i in range(0, len(data), 2):
        count = data[i]
        value = data[i + 1]
        decoded.extend([value] * count)

    return np.array(decoded).reshape(og_shape)

def count_repeats(data, start):
    value = data[start]
    count = 1

    for i in range(start + 1, len(data)):
        if data[i] != value or count == 128:
            break
        count += 1

    return count

def count_uniques(data, start):
    count = 1
    for i in range(start + 1, len(data)):
        if data[i] == data[i - 1] or count == 128:
            break
        count += 1

    return count

def byte_run_encode(img):
    flat = img.flatten()
    encoded = [img.shape[0], img.shape[1]]

    i = 0
    while i < len(flat):
        if i + 1 < len(flat) and flat[i] == flat[i + 1]:
            repeat_count = count_repeats(flat, i)
            while repeat_count > 128:
                encoded.append(-127)
                encoded.append(flat[i])
                repeat_count -= 128
            encoded.append(-(repeat_count - 1))
            encoded.append(flat[i])
            i += repeat_count

        else:
            unique_count = count_uniques(flat, i)
            while unique_count > 128:
                encoded.append(127)
                encoded.extend(flat[i:i+128])
                i += 128
                unique_count -= 128
            encoded.append(unique_count - 1)
            encoded.extend(flat[i:i+unique_count])
            i += unique_count

    # return img.shape, np.array(encoded)
    return np.array(encoded)

def byte_run_decode(encoded_data):
    og_shape = encoded_data[:2]
    encoded = encoded_data[2:]
    decoded = []

    i = 0
    while i < len(encoded):
        starting = encoded[i]
        if starting < 0:
            count = -starting + 1
            value = encoded[i + 1]
            decoded.extend([value] * count)
            i += 2

        else:
            count = starting + 1
            decoded.extend(encoded[i + 1:i + 1 + count])
            i += 1 + count

    return np.array(decoded).reshape(og_shape)

def compression_ratio(original, compressed):
    before_compression = get_size(original.flatten())
    after_compression = get_size(compressed)
    CR = before_compression / after_compression
    PR = after_compression / before_compression * 100

    return CR, PR

def get_size(obj, seen=None):
    import sys
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

def test():
    images = ['IMAGES/blueprint.png', 'IMAGES/document.png', 'IMAGES/mountains_panorama.png']

    for img_path in images:
        print(img_path)

        img = cv2.imread(img_path)
        img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY).astype(int)

        encoded = run_length_encode(img)
        decoded = run_length_decode(encoded)
        print("\tRLE: poprawność kompresji/dekompresji:", np.array_equal(img, decoded))
        CR, PR = compression_ratio(img, encoded[2:])
        print(f"\tRLE: stopień kompresji: {CR:.2f}, czyli {PR:.2f}%")

        encoded = byte_run_encode(img)
        decoded = byte_run_decode(encoded)
        print("\tByteRun: poprawność kompresji/dekompresji:", np.array_equal(img, decoded))
        CR, PR = compression_ratio(img, encoded[2:])
        print(f"\tByteRun: stopień kompresji: {CR:.2f}, czyli {PR:.2f}%\n")

def generate_report():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert
    import matplotlib.pyplot as plt

    document = Document()
    document.add_heading('Kompresja bezstratna\nWojciech Latos', 0)

    text = ('Zarządzanie pamięcią podczas kompresji jest bardzo ważne. Podczas testowania na przykładowych zdjęciach w każdym przypadku - oprócz ostatniego zdjęcia, przy metodzie RLE - udało się '
            'zmniejszyć rozmiar zajmowanej pamięci przez zdjęcie. Niepowodzenie wynikało z faktu, iż naturalne zdjęcie posiada bardzo duże spektrum barw, przez co kolory chociaż dla oka się nie zmieniają, '
            'to z komputerowego punktu widzenia już tak. Czarno-białe obrazy są idealne do kompresji tymi metodami, ponieważ są tylko dwie wartości - biały oraz czarny.\n'
            'Jednocześnie wedle samej nazwy - kompresja jest bezstratna - odkodowane obrazy są identyczne co oryginały.\n'
            'Sprawdzanie poprawności kompresji/dekompresji polega na porównaniu oryginalnego zdjęcia z tym po dekompresji '
            'przy pomocy kodu np.array_equal(oryginał, zakodowane) - zwraca True gdy tablice mają identyczne wartości oraz wymiary, False w przeciwnym wypadku.\n'
            'Oryginalny rozmiar obrazu znajduje się pod dwoma pierwszymi indeksami zakodowanej tablicy.')
    document.add_heading(text, level=1)

    imgs_path = ['IMAGES/blueprint.png', 'IMAGES/document.png', 'IMAGES/mountains_panorama.png']

    for img_path in imgs_path:
        print(img_path)
        img_imread = cv2.imread(img_path)
        img = cv2.cvtColor(img_imread, cv2.COLOR_RGB2GRAY).astype(int)

        plt.imshow(cv2.cvtColor(img_imread, cv2.COLOR_BGR2RGB))
        memfile = BytesIO()
        plt.savefig(memfile)
        memfile.seek(0)

        document.add_page_break()
        document.add_picture(memfile, width=Inches(6))
        plt.close()

        memfile.close()

        rle_encoded = run_length_encode(img)
        rle_decoded = run_length_decode(rle_encoded)
        document.add_paragraph('RLE: poprawność kompresji/dekompresji: {}'.format(np.array_equal(img, rle_decoded)))
        CR, PR = compression_ratio(img, rle_encoded[2:])
        document.add_paragraph('RLE: stopień kompresji: {:.2f}, czyli {:.2f}%'.format(CR, PR))

        br_encoded = byte_run_encode(img)
        br_decoded = byte_run_decode(br_encoded)
        document.add_paragraph('ByteRun: poprawność kompresji/dekompresji: {}'.format(np.array_equal(img, br_decoded)))
        CR, PR = compression_ratio(img, br_encoded[2:])
        document.add_paragraph('ByteRun: stopień kompresji: {:.2f}, czyli {:.2f}%'.format(CR, PR))

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')

# test()
generate_report()
