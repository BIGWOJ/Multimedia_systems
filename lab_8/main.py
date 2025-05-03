import matplotlib.pyplot as plt
import numpy as np
import cv2
import scipy.fftpack

class ver1:
    Y = np.array([])
    Cb = np.array([])
    Cr = np.array([])
    ChromaRatio = "4:4:4"
    QY = np.ones((8,8))
    QC = np.ones((8,8))
    shape = (0, 0, 3)

def dct2(a):
    return scipy.fftpack.dct(scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho')

def idct2(a):
    return scipy.fftpack.idct(scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

def zigzag(A):
    template= np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]

    return B

def CompressBlock(block, Q):
    block_centered = block - 128
    dct = dct2(block_centered)
    quantized = np.round(dct / Q).astype(int)

    return zigzag(quantized)

def DecompressBlock(vector, Q):
    block = zigzag(vector)
    dequantized = block * Q
    idct = idct2(dequantized)
    reconstructed = np.clip(idct + 128, 0, 255)

    return reconstructed.astype(np.uint8)

def CompressLayer(L, Q):
    S = np.array([])
    for w in range(0, L.shape[0], 8):
        for k in range(0, L.shape[1], 8):
            block = L[w:(w + 8), k:(k + 8)]
            S = np.append(S, CompressBlock(block, Q))

    return S

def DecompressLayer(S, Q, output_shape):
    height, width = output_shape
    L = np.zeros((height, width), dtype=np.uint8)
    blocks_w = width // 8

    for idx, i in enumerate(range(0, S.shape[0], 64)):
        vector = S[i:i+64]
        k = (idx % blocks_w) * 8
        w = (idx // blocks_w) * 8
        L[w:w+8, k:k+8] = DecompressBlock(vector, Q)

    return L

def CompressJPEG(RGB, Ratio="4:4:4", QY=np.ones((8,8)), QC=np.ones((8,8))):
    YCrCb = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
    Y = YCrCb[:, :, 0]
    Cb = YCrCb[:, :, 1].copy()
    Cr = YCrCb[:, :, 2].copy()

    JPEG = ver1()
    JPEG.shape = RGB.shape
    JPEG.ChromaRatio = Ratio

    # By default chroma subsampling is 4:4:4 (no subsampling)
    if Ratio == "4:2:2":
        Cb = Cb[:, ::2]
        Cr = Cr[:, ::2]

    JPEG.QY = QY
    JPEG.QC = QC
    JPEG.Y = CompressLayer(Y, QY)
    JPEG.Cb = CompressLayer(Cb, QC)
    JPEG.Cr = CompressLayer(Cr, QC)

    return JPEG

def DecompressJPEG(JPEG):
    Y_shape = JPEG.shape[:2]
    Y = DecompressLayer(JPEG.Y, JPEG.QY, Y_shape)

    if JPEG.ChromaRatio == "4:2:2":
        cb_shape = (Y_shape[0], Y_shape[1] // 2)
    else:
        cb_shape = Y_shape

    Cb = DecompressLayer(JPEG.Cb, JPEG.QC, cb_shape)
    Cr = DecompressLayer(JPEG.Cr, JPEG.QC, cb_shape)

    if JPEG.ChromaRatio == "4:2:2":
        Cb = np.repeat(Cb, 2, axis=1)
        Cr = np.repeat(Cr, 2, axis=1)

    YCrCb = np.stack([Y, Cb, Cr], axis=-1).clip(0, 255).astype(np.uint8)
    RGB = cv2.cvtColor(YCrCb, cv2.COLOR_YCrCb2RGB)
    return RGB

def run_length_encode(img):
    flat = np.array(img).flatten()
    encoded = [img.shape[0], img.shape[1]]
    count = 1

    for i in range(1, len(flat)):
        if flat[i] == flat[i - 1]:
            count += 1
        else:
            encoded.append(count)
            encoded.append(flat[i - 1])
            count = 1
    encoded.append(count)
    encoded.append(flat[-1])

    return np.array(encoded)

def test(QY, QC):
    img = cv2.imread("images/img_2.jpg")
    img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    x_start = np.random.randint(0, img_rgb.shape[1] - 128)
    y_start = np.random.randint(0, img_rgb.shape[0] - 128)
    fragment = img_rgb[y_start:y_start + 128, x_start:x_start + 128]

    fragment_compressed = CompressJPEG(fragment, Ratio="4:2:2", QY=QY, QC=QC)
    fragment_decompressed = DecompressJPEG(fragment_compressed)

    fig, axs = plt.subplots(4, 2, sharey=True)
    fig.suptitle("Porównanie oryginalnych warstw z warstwami po dekompresji")
    fig.set_size_inches(9,13)

    axs[0,0].imshow(fragment)
    axs[0,0].set_title("Oryginalny obraz")

    before_YCrCb=cv2.cvtColor(fragment,cv2.COLOR_RGB2YCrCb)
    axs[1,0].set_title("Warstwa Y")
    axs[1,0].imshow(before_YCrCb[:,:,0],cmap=plt.cm.gray)
    axs[2,0].set_title("Warstwa Cb")
    axs[2,0].imshow(before_YCrCb[:,:,1],cmap=plt.cm.gray)
    axs[3,0].set_title("Warstwa Cr")
    axs[3,0].imshow(before_YCrCb[:,:,2],cmap=plt.cm.gray)

    axs[0,1].set_title("Obraz po dekompresji")
    axs[0,1].imshow(fragment_decompressed)
    after_YCrCb=cv2.cvtColor(fragment_decompressed,cv2.COLOR_RGB2YCrCb)
    axs[1,1].set_title("Warstwa Y")
    axs[1,1].imshow(after_YCrCb[:,:,0],cmap=plt.cm.gray)
    axs[2,1].set_title("Warstwa Cb")
    axs[2,1].imshow(after_YCrCb[:,:,1],cmap=plt.cm.gray)
    axs[3,1].set_title("Warstwa Cr")
    axs[3,1].imshow(after_YCrCb[:,:,2],cmap=plt.cm.gray)

    plt.show()

def generate_report(QY, QC):
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert
    from os import remove
    import matplotlib.pyplot as plt

    files_path = ['images/img_1.jpg', 'images/img_2.jpg', 'images/img_3.jpg', 'images/img_4.jpg']
    chroma_ratios = ['4:2:2', '4:4:4']

    document = Document()
    document.add_heading('Implementacja częściowej kompresji JPEG\nWojciech Latos', 0)
    document.add_paragraph('W tytule zawarte są parametry, dla których został wykonany test algorytmu kompresji.\n'
                           'Wartość True dla QY/QC oznacza przekazania do algorytmu konkretnej tablicy kwantyzacji, natomiast False - użycie domyślnej macierzy jedynek.\n'
                           'W celu uzyskania jak najlepszych wyników, fragmenty obrazu zostały wybrane losowo.')
    document.add_paragraph('Poniżej zostały przedstawione wyniki kompresji częściowej kompresji JPEG z różnymi parametrami, na konkretnych warstwach.\n'
                           'Pod fragmentem danej warstwy jest podana liczba określona w procentach jako stosunek długości skompresowanego wektora warstwy względem adekwatnej warstwy oryginalnego obrazu.\n'
                           'Jak widać, mimo faktu, iż wektor "skompresowany" warstwy Y jest dłuższy od oryginalnego, to już warstwy Cb oraz Cr są znacząco mniejsze. '
                           'Finalnie nakładając wszystkie 3 warstwy na siebie otrzymujemy obraz jakościowo nieodbiegający od oryginału.\n'
                           'Uwzględniając długości wektorów warstw fragmentów obrazów finalny rozmiar skompresowanej wersji jest mniejszy względem oryginału.\n'
                           'Najgorszymi wynikami okazał się wariant ratio 4:4:4 oraz QY, QC jako macierz jedynek.\n'
                           'W celu określenia procentowego skrócenia dla warstw użyty został algorytm RLE.')

    for file_path in files_path:
        print('file')
        img = cv2.imread(file_path)
        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        plt.imshow(img_rgb)
        plt.title("Oryginalny obraz")

        memfile = BytesIO()
        plt.savefig(memfile)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        plt.close()

        for _ in range(4):
            x_start = np.random.randint(0, img_rgb.shape[1] - 128)
            y_start = np.random.randint(0, img_rgb.shape[0] - 128)
            fragment = img_rgb[y_start:y_start+128, x_start:x_start+128]
            for QY_matrix, QC_matrix in [(QY, QC), (None, None)]:
                for ratio in chroma_ratios:
                    if QY_matrix is None:
                        fragment_compressed = CompressJPEG(fragment, Ratio=ratio)
                    else:
                        fragment_compressed = CompressJPEG(fragment, Ratio=ratio, QY=QY_matrix, QC=QC_matrix)
                    fragment_decompressed = DecompressJPEG(fragment_compressed)

                    fig, axs = plt.subplots(4, 2, sharey=True, gridspec_kw={'hspace': 0.5})
                    fig.suptitle(f"Porównanie oryginalnych warstw z warstwami po dekompresji\n"
                                 f"Ratio: {ratio}, QY: {QY_matrix is not None}, QC: {QC_matrix is not None}", fontsize=16)
                    fig.set_size_inches(9, 13)

                    axs[0, 0].imshow(fragment)
                    axs[0, 0].set_title("Oryginalny fragment")
                    before_YCrCb = cv2.cvtColor(fragment, cv2.COLOR_RGB2YCrCb)
                    axs[1, 0].set_title("Warstwa Y")
                    axs[1, 0].imshow(before_YCrCb[:, :, 0], cmap=plt.cm.gray)
                    axs[2, 0].set_title("Warstwa Cb")
                    axs[2, 0].imshow(before_YCrCb[:, :, 1], cmap=plt.cm.gray)
                    axs[3, 0].set_title("Warstwa Cr")
                    axs[3, 0].imshow(before_YCrCb[:, :, 2], cmap=plt.cm.gray)
                    axs[0, 1].set_title("Fragment po dekompresji")
                    axs[0, 1].imshow(fragment_decompressed)

                    after_YCrCb = cv2.cvtColor(fragment_decompressed, cv2.COLOR_RGB2YCrCb)
                    axs[1, 1].set_title(f"Warstwa Y\n"
                                        f"{int(run_length_encode(after_YCrCb[:, :, 0]).shape[0] / 16384 * 100)}% względem oryginału")
                    axs[1, 1].imshow(after_YCrCb[:, :, 0], cmap=plt.cm.gray)
                    axs[2, 1].set_title(f"Warstwa Cb\n"
                                        f"{int(run_length_encode(after_YCrCb[:, :, 1]).shape[0] / 16384 * 100)}% względem oryginału")
                    axs[2, 1].imshow(after_YCrCb[:, :, 1], cmap=plt.cm.gray)
                    axs[3, 1].set_title(f"Warstwa Cr\n"
                                        f"{int(run_length_encode(after_YCrCb[:, :, 2]).shape[0] / 16384 * 100)}% względem oryginału")
                    axs[3, 1].imshow(after_YCrCb[:, :, 2], cmap=plt.cm.gray)

                    fig.tight_layout()
                    memfile = BytesIO()
                    fig.savefig(memfile)
                    memfile.seek(0)
                    document.add_picture(memfile, width=Inches(6))
                    memfile.close()
                    plt.close(fig)

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')



QY = np.array([
        [16, 11, 10, 16, 24, 40, 51, 61],
        [12, 12, 14, 19, 26, 58, 60, 55],
        [14, 13, 16, 24, 40, 57, 69, 56],
        [14, 17, 22, 29, 51, 87, 80, 62],
        [18, 22, 37, 56, 68, 109, 103, 77],
        [24, 36, 55, 64, 81, 104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
    ])

QC = np.array([
    [17, 18, 24, 47, 99, 99, 99, 99],
    [18, 21, 26, 66, 99, 99, 99, 99],
    [24, 26, 56, 99, 99, 99, 99, 99],
    [47, 66, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
])

# test(QY=QY, QC=QC)
generate_report(QY=QY, QC=QC)
