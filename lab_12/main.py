import numpy as np
import cv2
from skimage.metrics import structural_similarity
import matplotlib.pyplot as plt

def draw_line(image, p1, p2, color):
    x1, y1 = map(int, p1)
    x2, y2 = map(int, p2)
    dx = abs(x2 - x1)
    dy = abs(y2 - y1)
    sx = 1 if x2 > x1 else -1
    sy = 1 if y2 > y1 else -1
    err = dx - dy

    while True:
        if 0 <= x1 < image.shape[1] and 0 <= y1 < image.shape[0]:
            image[y1, x1] = color
        if x1 == x2 and y1 == y2:
            break
        e2 = 2 * err
        if e2 > -dy:
            err -= dy
            x1 += sx
        if e2 < dx:
            err += dx
            y1 += sy

def fill_polygon(image, points, color):
    pts = np.array(points, np.int32)
    pts = pts.reshape((-1, 1, 2))
    cv2.fillPoly(image, [pts], color=color)

def draw_circle(image, center, radius, color):
    cx, cy = map(int, center)
    x = 0
    y = radius
    p = 1 - radius

    border_points = []

    while x <= y:
        points = [
            (cx + x, cy + y), (cx + y, cy + x),
            (cx - x, cy + y), (cx - y, cy + x),
            (cx + x, cy - y), (cx + y, cy - x),
            (cx - x, cy - y), (cx - y, cy - x)
        ]
        for px, py in points:
            if 0 <= px < image.shape[1] and 0 <= py < image.shape[0]:
                image[py, px] = color
                border_points.append([px, py])

        x += 1
        if p < 0:
            p += 2 * x + 1
        else:
            y -= 1
            p += 2 * (x - y) + 1

def draw_triangle(image, p1, p2, p3, color, filled=False):
    if filled:
        fill_polygon(image, [p1, p2, p3], color)
    else:
        draw_line(image, p1, p2, color)
        draw_line(image, p2, p3, color)
        draw_line(image, p3, p1, color)

def draw_rectangle(image, p1, p2, color, filled=False):
    x1, y1 = map(int, p1)
    x2, y2 = map(int, p2)

    if filled:
        fill_polygon(image, [[x1, y1], [x2, y1], [x2, y2], [x1, y2]], color)
    else:
        draw_line(image, (x1, y1), (x2, y1), color)
        draw_line(image, (x2, y1), (x2, y2), color)
        draw_line(image, (x2, y2), (x1, y2), color)
        draw_line(image, (x1, y2), (x1, y1), color)

def generate_image(canvas_width, canvas_height, background_color, shapes):
    image = np.full((canvas_height, canvas_width, 3), background_color, dtype=np.uint8)
    shapes.sort(key=lambda s: s.get("Z_layer", 0))

    for shape in shapes:
        t = shape["type"]
        c = shape["color"]
        filled = shape.get("filled", False)
        if t == "line":
            draw_line(image, shape["p1"], shape["p2"], c)
        elif t == "circle":
            draw_circle(image, shape["center"], shape["radius"], c)
        elif t == "rectangle":
            draw_rectangle(image, shape["p1"], shape["p2"], c, filled=filled)
        elif t == "triangle":
            draw_triangle(image, shape["p1"], shape["p2"], shape["p3"], c, filled=filled)
        elif t == "polygon":
            if filled:
                fill_polygon(image, shape["points"], c)
            else:
                points = shape["points"]
                for i in range(len(points)):
                    draw_line(image, points[i], points[(i + 1) % len(points)], c)

    return image

def compare_images(img1, img2):
    mse = np.mean((img1 - img2) ** 2)
    ssim = structural_similarity(img1, img2, multichannel=True, channel_axis=2, win_size=3)
    return mse, ssim

def generate_report():
    from docx import Document
    from docx2pdf import convert
    from os import remove

    def add_image_to_docx(document, image, title):
        from io import BytesIO
        from matplotlib.figure import Figure
        from docx.shared import Inches

        memfile = BytesIO()
        if isinstance(image, Figure) or isinstance(image, plt.Axes):
            image.savefig(memfile)
            plt.close(image)
        else:
            fig = plt.figure()
            plt.imshow(image)
            plt.title(title)
            fig.savefig(memfile)
            plt.close(fig)
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(7))
        memfile.close()

    document = Document()
    document.add_heading('Rasteryzacja danych wektorowych\nWojciech Latos', 0)
    document.add_paragraph("Wygenerowany obraz jest zgodny z założonymi wymaganiami. Ostatnią cyfrą mojego numer indeksu (53894) jest 4. "
                           "Dlatego dodatkowo został wygenerowany trójkąt, prostokąt oraz okrąg (prostokąt z tyłu, trójkąt częściowo przysłania go, na wierzchu okrąg).")

    shapes = [
        {"type": "circle", "Z_layer": 0, "center": [100, 100], "radius": 60, "color": [0, 255, 0]},
        {"type": "triangle", "Z_layer": 1, "p1": [60, 65], "p2": [140, 65], "p3": [100, 15], "color": [0, 255, 255],
         "filled": True},

        {"type": "rectangle", "Z_layer": 0, "p1": [200, 50], "p2": [350, 150], "color": [255, 255, 0]},
        {"type": "rectangle", "Z_layer": 1, "p1": [220, 70], "p2": [240, 90], "color": [0, 0, 255], "filled": True},
        {"type": "rectangle", "Z_layer": 1, "p1": [260, 110], "p2": [280, 130], "color": [255, 0, 0], "filled": True},

        {
            "type": "polygon",
            "Z_layer": 0,
            "points": [
                [400, 200],
                [420, 200],
                [420, 300],
                [460, 300],
                [460, 320],
                [400, 320],
                [400, 200]
            ],
            "color": [255, 0, 255]
        },

        {"type": "circle", "Z_layer": 0, "center": [600, 150], "radius": 40, "color": [255, 255, 0]},
        {"type": "rectangle", "Z_layer": 1, "p1": [500, 100], "p2": [600, 200], "color": [100, 70, 0], "filled": True},

        {"type": "rectangle", "Z_layer": 0, "p1": [100, 300], "p2": [300, 400], "color": [255, 0, 0], "filled": True},
        {"type": "triangle", "Z_layer": 1, "p1": [120, 350], "p2": [280, 350], "p3": [200, 250], "color": [255, 255, 0],
         "filled": True},
        {"type": "circle", "Z_layer": 2, "center": [200, 350], "radius": 30, "color": [0, 0, 255]},
    ]

    image_sizes = [
        (800, 600),
        (950, 700),
        (1100, 800),
        (1250, 900),
        (1400, 1000),
    ]
    background_color = [0, 0, 0]

    generated_images = []
    for size in image_sizes:
        img = generate_image(size[0], size[1], background_color, shapes)
        generated_images.append(img)

    target_size = (650, 500)
    resized_images = []

    for idx, img in enumerate(generated_images):
        resized_img = cv2.resize(img, target_size)
        resized_images.append(resized_img)

    base_image = generate_image(target_size[0], target_size[1], background_color, shapes)
    add_image_to_docx(document, base_image, "Wygenerowany obraz")

    document.add_page_break()
    document.add_paragraph("Na podstawie poniższego zestawienia można zauważyć, że obrazy o różnych rozmiarach, które zostały przeskalowane do identycznego rozmiaru (650x500) różnią się wartościami MSE i SSIM. "
                           "Im proces skalowania był mocniejszy (większy rozmiar startowy) tym wartość MSE była niższa, natomiast SSIM było większe.")

    fig = plt.figure(figsize=(15, 5))
    for i in range(5):
        metrics = compare_images(base_image, resized_images[i])
        plt.subplot(1, 5, i + 1)
        plt.imshow(resized_images[i])
        plt.title(f"Obraz {image_sizes[i][0]} x {image_sizes[i][1]}\n"
                  f"MSE={metrics[0]:.2f}, SSIM={metrics[1]:.4f}")
        plt.axis('off')
    add_image_to_docx(document, fig, "Porównanie obrazów")

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')
    remove(docx_path)

generate_report()
