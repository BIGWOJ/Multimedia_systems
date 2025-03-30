import matplotlib.pyplot as plt
import numpy as np
import cv2

def color_fit(pixel_color, color_palette):
    distances = np.linalg.norm(color_palette - pixel_color, axis=1)
    return color_palette[np.argmin(distances)]

def quantize_image(img, color_palette):
    if len(color_palette.shape) == 1:
        color_palette = np.stack([color_palette] * 3, axis=1)

    img_normalized = img / 255.0
    out_img = np.zeros_like(img_normalized)
    for row in range(img_normalized.shape[0]):
        for col in range(img_normalized.shape[1]):
            out_img[row, col] = color_fit(img_normalized[row, col], color_palette)

    return (out_img * 255).astype(np.uint8)

def dither_random_binary(img):
    if len(img.shape) == 3:
        img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

    img = img / 255.0
    random_matrix = np.random.rand(img.shape[0], img.shape[1])
    dithered_img = img >= random_matrix
    dithered_img = dithered_img * 1

    return dithered_img

def dither_ordered(img, color_palette):
    if len(color_palette.shape) == 1:
        color_palette = np.stack([color_palette] * 3, axis=1)

    threshold_map_m2 = np.array([[0,8,2,10],
                                 [12,4,14,6],
                                 [3,11,1,9],
                                 [15,7,13,5]])
    r = 1
    M = threshold_map_m2
    n = int(threshold_map_m2.shape[0] / 2)
    Mpre = (M+1) / (2*n)**2 - 0.5
    img_normalized = img / 255.0

    out_img = np.zeros_like(img_normalized)
    for row in range(img_normalized.shape[0]):
        for col in range(img_normalized.shape[1]):
            mpre_value = Mpre[row % (2*n), col % (2*n)]
            out_img[row, col] = color_fit(img_normalized[row, col] + r * mpre_value, color_palette)

    return (out_img * 255).astype(np.uint8)

def dither_floyd_steinberg(img, color_palette):
    if len(color_palette.shape) == 1:
        color_palette = np.stack([color_palette] * 3, axis=1)

    img_normalized = img / 255.0
    out_img = np.zeros_like(img_normalized)
    for row in range(img_normalized.shape[0]):
        for col in range(img_normalized.shape[1]):
            old_pixel = img_normalized[row, col]
            new_pixel = color_fit(old_pixel, color_palette)
            out_img[row, col] = new_pixel
            quant_error = old_pixel - new_pixel

            if row < img_normalized.shape[0] - 1:
                img_normalized[row + 1, col] += quant_error * 7 / 16
            if col > 0 and row < img_normalized.shape[0] - 1:
                img_normalized[row + 1, col - 1] += quant_error * 3 / 16
            if col < img_normalized.shape[1] - 1:
                img_normalized[row, col + 1] += quant_error * 5 / 16
            if col < img_normalized.shape[1] - 1 and row < img_normalized.shape[0] - 1:
                img_normalized[row + 1, col + 1] += quant_error * 1 / 16

    return (out_img * 255).astype(np.uint8)

def test_quantize(img, color_palette, generate_report=False):
    img_quantized = quantize_image(img, color_palette)

    if generate_report:
        return img_quantized

    plt.imshow(img_quantized)
    plt.show()

def test_dither(img, method, color_palette=None, generate_report=False):
    if method == 'random_binary':
        dithered_img = dither_random_binary(img)
    elif method == 'ordered':
        dithered_img = dither_ordered(img, color_palette)
    elif method == 'floyd_steinberg':
        dithered_img = dither_floyd_steinberg(img, color_palette)

    if generate_report:
        return dithered_img

    plt.imshow(dithered_img, cmap='gray')
    plt.show()

def generate_comparisions(img, color_palette, generate_report=False):
    fig = plt.figure(figsize=(10, 7))

    if len(color_palette.shape) == 2:
        fig.suptitle(f'Dithering paleta {color_palette.shape[0]} kolorów', fontsize=16)
    else:
        if int(np.sqrt(color_palette.shape[0])) == 1:
            fig.suptitle(f'Dithering {int(np.sqrt(color_palette.shape[0]))}-bit', fontsize=16)
        else:
            fig.suptitle(f'Dithering {int(np.sqrt(color_palette.shape[0]))}-bity', fontsize=16)

    # Adding random dithering for 1-bit color palette
    subplot_index_change = 0
    if color_palette.shape[0] == 2:
        subplot_index_change = 1
        plt.subplot(2, 3, 2)
        plt.title("Dithering losowy")
        img_dithered = test_dither(img, 'random_binary', color_palette, True)
        plt.imshow(img_dithered, cmap='gray')

    plt.subplot(2, 2+subplot_index_change, 1)
    plt.title("Oryginał")
    plt.imshow(img)

    plt.subplot(2, 2+subplot_index_change, 2+subplot_index_change)
    plt.title("Dithering zorganizowany")
    img_dithered = test_dither(img, 'ordered', color_palette, True)
    plt.imshow(img_dithered, cmap='gray')

    plt.subplot(2, 2+subplot_index_change, 3+subplot_index_change)
    plt.title("Kwantyzacja")
    img_quantized = test_quantize(img, color_palette, True)
    plt.imshow(img_quantized)

    plt.subplot(2, 2+subplot_index_change, 4+subplot_index_change)
    plt.title("Dithering Floyda-Steinberga")
    img_dithered = test_dither(img, 'floyd_steinberg', color_palette, True)
    plt.imshow(img_dithered, cmap='gray')

    plt.tight_layout()

    if generate_report:
        return fig

    plt.show()

def tests():
    pallet8 = np.array([
        [0.0, 0.0, 0.0, ],
        [0.0, 0.0, 1.0, ],
        [0.0, 1.0, 0.0, ],
        [0.0, 1.0, 1.0, ],
        [1.0, 0.0, 0.0, ],
        [1.0, 0.0, 1.0, ],
        [1.0, 1.0, 0.0, ],
        [1.0, 1.0, 1.0, ],
    ])

    pallet16 = np.array([
        [0.0, 0.0, 0.0, ],
        [0.0, 1.0, 1.0, ],
        [0.0, 0.0, 1.0, ],
        [1.0, 0.0, 1.0, ],
        [0.0, 0.5, 0.0, ],
        [0.5, 0.5, 0.5, ],
        [0.0, 1.0, 0.0, ],
        [0.5, 0.0, 0.0, ],
        [0.0, 0.0, 0.5, ],
        [0.5, 0.5, 0.0, ],
        [0.5, 0.0, 0.5, ],
        [1.0, 0.0, 0.0, ],
        [0.75, 0.75, 0.75, ],
        [0.0, 0.5, 0.5, ],
        [1.0, 1.0, 1.0, ],
        [1.0, 1.0, 0.0, ]
    ])

    pallet_gray = np.linspace(0, 1, 2)

    img_small = cv2.imread('IMG_GS/GS_0001.tif')
    img_small = cv2.cvtColor(img_small, cv2.COLOR_BGR2RGB)

    plt.imshow(img_small)
    plt.show()

    test_quantize(img_small, pallet_gray)
    test_quantize(img_small, pallet_gray)
    test_quantize(img_small, pallet_gray)
    test_dither(img_small, 'random_binary')
    test_dither(img_small, 'ordered', pallet_gray)
    test_dither(img_small, 'floyd_steinberg', pallet_gray)
    generate_comparisions(img_small, pallet_gray)

def generate_report():
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    from docx2pdf import convert
    from os import remove

    imgs_grayscale = ['IMG_GS/GS_0001.tif', 'IMG_GS/GS_0002.png', 'IMG_GS/GS_0003.png']
    imgs_rgb = ['IMG_SMALL/SMALL_0004.jpg', 'IMG_SMALL/SMALL_0006.jpg', 'IMG_SMALL/SMALL_0007.jpg', 'IMG_SMALL/SMALL_0009.jpg']
    files = imgs_grayscale + imgs_rgb
    quantize_range_bits = [1, 2, 4]
    pallet8 = np.array([
        [0.0, 0.0, 0.0, ],
        [0.0, 0.0, 1.0, ],
        [0.0, 1.0, 0.0, ],
        [0.0, 1.0, 1.0, ],
        [1.0, 0.0, 0.0, ],
        [1.0, 0.0, 1.0, ],
        [1.0, 1.0, 0.0, ],
        [1.0, 1.0, 1.0, ],
    ])
    pallet16 = np.array([
        [0.0, 0.0, 0.0, ],
        [0.0, 1.0, 1.0, ],
        [0.0, 0.0, 1.0, ],
        [1.0, 0.0, 1.0, ],
        [0.0, 0.5, 0.0, ],
        [0.5, 0.5, 0.5, ],
        [0.0, 1.0, 0.0, ],
        [0.5, 0.0, 0.0, ],
        [0.0, 0.0, 0.5, ],
        [0.5, 0.5, 0.0, ],
        [0.5, 0.0, 0.5, ],
        [1.0, 0.0, 0.0, ],
        [0.75, 0.75, 0.75, ],
        [0.0, 0.5, 0.5, ],
        [1.0, 1.0, 1.0, ],
        [1.0, 1.0, 0.0, ]
    ])
    color_pallets = [pallet8, pallet16]

    document = Document()
    document.add_heading('Kwantyzacja obrazu i dithering\nWojciech Latos', 0)

    for file in files:
        img = cv2.imread(file)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        if file.__contains__('GS'):
            dithering_pallets = [np.linspace(0, 1, 2 ** bit) for bit in quantize_range_bits]

        else:
            dithering_pallets = color_pallets

        for dithering_pallet in dithering_pallets:
            comparison_fig = generate_comparisions(img, dithering_pallet, True)
            comparison_fig_memfile = BytesIO()
            comparison_fig.savefig(comparison_fig_memfile, format='png')
            comparison_fig_memfile.seek(0)

            document.add_picture(comparison_fig_memfile, width=Inches(7))

            comparison_fig_memfile.close()
            plt.close(comparison_fig)

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')
    remove(docx_path)

# tests()
generate_report()
