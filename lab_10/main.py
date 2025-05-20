import os
import cv2
import numpy as np
import matplotlib.pyplot as plt
from skimage.metrics import structural_similarity


def compute_metrics(original, modified):
    original = cv2.cvtColor(original, cv2.COLOR_RGB2GRAY)
    modified = cv2.cvtColor(modified, cv2.COLOR_RGB2GRAY)

    mse = np.mean((original - modified) ** 2)
    nmse = mse / (np.mean(modified ** 2))
    psnr = 10 * np.log10((255 ** 2) / mse)
    if_rating = np.sum((modified - original) ** 2) / (np.sum(modified * original))
    if_ssim = structural_similarity(original, modified, data_range=modified.max() - modified.min())

    return {
        "MSE": mse,
        "NMSE": nmse,
        "PSNR": psnr,
        "IF": if_rating,
        "SSIM": if_ssim
    }

def get_metrics_text(metrics):
    return f"MSE: {metrics['MSE']:.2f} NMSE: {metrics['NMSE']:.2f} PSNR: {metrics['PSNR']:.2f}\n" \
           f"IF: {metrics['IF']:.2f} SSIM: {metrics['SSIM']:.2f}"


def compress_jpeg(img, quality):
    encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), quality]
    result, encimg = cv2.imencode('.jpg', img, encode_param)
    decimg = cv2.imdecode(encimg, 1)
    return decimg


def blur_image(img, kernel_size):
    return cv2.blur(img, (kernel_size, kernel_size))


def gaussian_noise_image(img, alpha=0.5, sigma=25):
    gauss = np.random.normal(0, sigma, img.shape).astype(np.int16)
    noisy = (img + alpha * gauss).clip(0,255).astype(np.uint8)
    return noisy

def add_image_to_docx(document, image, title):
    from io import BytesIO
    from docx.shared import Inches

    fig = plt.figure()
    plt.imshow(image)
    plt.title(title)
    memfile = BytesIO()
    fig.savefig(memfile)
    memfile.seek(0)
    document.add_picture(memfile, width=Inches(6))
    memfile.close()
    plt.close(fig)

def generate_report():
    from docx import Document
    from docx2pdf import convert
    from os import remove
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.add_heading('Ocena jakości obrazów — miary obiektywne\nWojciech Latos', 0)
    document.add_heading('Wnioski można przeczytać na końcu sprawozdania', level=1)

    images_dir = 'images'
    for image_counter, image_file in enumerate(os.listdir(images_dir)):
        original_img = cv2.imread(os.path.join(images_dir, image_file))
        original_img = cv2.cvtColor(original_img, cv2.COLOR_BGR2RGB)
        add_image_to_docx(document, original_img, 'Obraz oryginalny')

        if image_counter == 0:
            jpeg_metrics = []
            for jpeg_ratio in [10, 20, 30, 40, 50, 60, 75]:
                modified_img = compress_jpeg(original_img, jpeg_ratio)
                metrics = compute_metrics(original_img, modified_img)
                jpeg_metrics.extend(metrics.values())
                add_image_to_docx(document, modified_img, f'Kompresja JPEG {jpeg_ratio}\n\n{get_metrics_text(metrics)}')

        elif image_counter == 1:
            blur_metrics = []
            for blur in [2, 4, 6, 8, 10, 12, 14]:
                modified_img = blur_image(original_img, blur)
                metrics = compute_metrics(original_img, modified_img)
                blur_metrics.extend(metrics.values())
                add_image_to_docx(document, modified_img, f'Filtracja uśredniająca {blur}\n\n{get_metrics_text(metrics)}')

        elif image_counter == 2:
            gauss_metrics = []
            for gauss in [0.1, 0.25, 0.4, 0.55, 0.7, 0.85, 1.0]:
                modified_img = gaussian_noise_image(original_img, alpha=gauss)
                metrics = compute_metrics(original_img, modified_img)
                gauss_metrics.extend(metrics.values())
                add_image_to_docx(document, modified_img, f'Zaszumienie obrazu {gauss}\n\n{get_metrics_text(metrics)}')

    document.add_page_break()
    document.add_heading('Tabela zbiorcza wyników', level=1)
    tables = [
        document.add_table(rows=5, cols=8),
        document.add_table(rows=5, cols=8),
        document.add_table(rows=5, cols=8),
        document.add_table(rows=5, cols=8),
        document.add_table(rows=5, cols=8),
    ]

    methods_list = ['Kompresja JPEG (10, 20, 30, 40, 50, 60, 75)',
                    'Filtracja uśredniająca / blur (2, 4, 6, 8, 10, 12, 14)',
                    'Zaszumienie obrazu Gaussowskie (0.1, 0.25, 0.4, 0.55, 0.7, 0.85, 1)']
    parameters_list = ['Parametr 1', 'Parametr 2', 'Parametr 3', 'Parametr 4', 'Parametr 5', 'Parametr 6', 'Parametr 7']
    method_metrics_list = [jpeg_metrics, blur_metrics, gauss_metrics]
    headers = ['MSE', 'NMSE', 'PSNR', 'IF', 'SSIM']

    for counter, (table, header) in enumerate(zip(tables, headers)):
        table.style = 'Table Grid'
        top_row = table.rows[0]
        merged_cell = top_row.cells[0].merge(top_row.cells[-1])
        merged_cell.text = header
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        second_row = table.rows[1]
        for i, parameter in enumerate(parameters_list):
            second_row.cells[i + 1].text = parameter
            second_row.cells[i + 1].paragraphs[0].runs[0].font.size = Pt(10)

        for method_counter, metrics in enumerate(method_metrics_list):
            table.rows[2 + method_counter].cells[0].text = methods_list[
                method_counter]

            for i, metric in enumerate(metrics[counter::5]):
                  table.rows[2 + method_counter].cells[i + 1].text = str(np.round(metric, 2))

    document.add_heading('Wskaźnik MSE', level=2)
    document.add_paragraph('Podczas kompresji JPEG zmniejszał się wraz ze wzrostem jakości kompresji (wyższa wartość argumentu).'
                           'Wartość dla ostatniego argumentu to tylko 40% wartości pierwszego. Przez co widać dużą zmianę.\n'
                           'Podczas filtracji uśredniającej wskaźnik tym razem wzrastał wraz ze wzrostem argumentu filtracji.'
                           'Tym razem różnice między wartościami dla paremtrów nie były tak duże jak w przypadku kompresji. Największa wartość była ponad 75% wartości największej.\n'
                           'Największe różnice zanotowane zostały podczas zaszuemiania obrazu. Początkowa wartość 3.05 wzrostła wraz ze wzrostem wartości zaszumiania aż do 80.59. Co przekłada się na wzrost ponad 26-krotny.')

    document.add_heading('Wskaźnik NMSE', level=2)
    document.add_paragraph('Podczas kompresji JPEG wartości zmniejszały się stabilnie, w równym tempie. Startowe 0.36 zakończyło się na 0.13 co przekłada się na prawie 30% wartości początkowej.\n'
                           'Podczas filtracji wartości wzrastały także w równym tempie. Najmniejsza wartość wyniosła 80% tej największej.\n'
                           'Największe wzrosty zostały znowu zauważone podczas zaszumiania obrazu. Początkowa wartość 0.03 wzrostła aż do 0.72. Przekłada się to na wzrost 24-krotny.')

    document.add_page_break()

    document.add_heading('Wskaźnik PSNR', level=2)
    document.add_paragraph('Podczas kompresji JPEG wartości minimalnie się zwiększały. Startowa wartość 32.46 urosła tylko do 36.4. Jest to bardzo mały wzrost.\n'
                           'Podczas filtracji tym razem wartości zmniejszały się. Odbywało się to w powolnym tempie (zakres to 30.44 - 29.37).\n'
                           'Zaszumienie obrazu po raz kolejny zaowocowało największymi wzrostami. Startowa wartośći 43.28 urosła do 29.07. Procentowo wzrosty nie były aż tak duże jak w poprzednich badaniach natomiast na tle wskaźnika PSNR były największe.')

    document.add_heading('Wskaźnik IF', level=2)
    document.add_paragraph('Wskaźnik podczas testowania kompresji malał liniowo. Zakres to 0.3 do 0.12.\n'
                           'Wartości filtracji uśredniającej tym razem wzrastały wraz ze wzrostem argumentu. Wzrost był mały, stabilny, ciągły, liniowy.\n'
                           'Testowanie zaszumienia ponownie skutkowało największymi zmianami w wartościach. Rekordoy wzrost o 32 krotność.')

    document.add_heading('Wskaźnik SSIM', level=2)
    document.add_paragraph('Kompresja JPEG tym razem wzrastała. Wzrost stabilny, liniowy, początkowe 0.79 wzrostło do 0.94\n'
                           'Filtracja w przypadku wskaźnika SSIM spadała wraz ze wzrostem argumentu. Końcowa wartość to było 50% początkowej 0.74\n'
                           'Tym razem zaszumienie obrazu nie przyniosło spektakularnych zmian. Wartości malały stabilnie, liniowo. Najmniejsza wartość to 55% początkowej. Najmniejsze skoki wartości względem wszystkich wskaźników.')

    document.add_heading('Podsumowanie', level=2)
    document.add_paragraph('Zmiana wskaźników podczas kompresji JPEG cechowała się liniową zależnością, powolną, stabilną zmianą. Niektóre wskaźniki zmniejszały się inne wzrastały wraz ze wzrostem modyfikacji obrazu.\n'
                           'Filtracja uśredniająca także analogicznie do kompresji. Wartości zachowywały się bardzo podobnie.\n'
                           'Podczas zaszumiania obrazów wartości zmieniały się w największym zakresie ze wszystkich modyfikacji. Silnie się zmieniały (zmniejszały/zwiększały) wraz ze wzrostami argumentów.')

    docx_path = 'report.docx'
    document.save(docx_path)
    convert(docx_path, 'report.pdf')
    remove(docx_path)


generate_report()